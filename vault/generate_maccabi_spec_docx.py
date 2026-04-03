#!/opt/homebrew/bin/python3
"""
Generate Hebrew RTL Word document for Maccabi Clicks Integration Spec.
Uses python-docx with proper RTL layout, Hebrew fonts, and embedded screenshots.
"""

import os
import io
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image as PILImage

SCREENSHOTS_DIR = Path('/Users/ramikhouri/Desktop/Taliaz/maccabi-spec-screenshots')
OUTPUT_DOCX = Path('/Users/ramikhouri/Desktop/Taliaz/Maccabi_Clicks_Integration_Spec.docx')

# Colors
DARK_BLUE = RGBColor(0x0A, 0x3D, 0x62)
BRAND_BLUE = RGBColor(0x00, 0x77, 0xB6)
RED = RGBColor(0xC0, 0x39, 0x2B)
DARK_GRAY = RGBColor(0x1A, 0x1A, 0x1A)
MEDIUM_GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x27, 0xAE, 0x60)
PURPLE = RGBColor(0x7B, 0x2D, 0x8E)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
YELLOW = RGBColor(0xF3, 0x9C, 0x12)


def set_rtl_paragraph(paragraph):
    """Set paragraph to RTL direction."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>')
    pPr.append(bidi)


def set_rtl_run(run):
    """Set run to RTL direction."""
    rPr = run._r.get_or_add_rPr()
    rtl = parse_xml(f'<w:rtl {nsdecls("w")} val="1"/>')
    rPr.append(rtl)


def set_cell_rtl(cell):
    """Set table cell to RTL."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # No direct RTL on tcPr, we set it on paragraphs inside
    for p in cell.paragraphs:
        set_rtl_paragraph(p)


def add_hebrew_run(paragraph, text, bold=False, italic=False, size=Pt(11), color=DARK_GRAY, font_name='Arial'):
    """Add a Hebrew text run with proper RTL settings."""
    run = paragraph.add_run(text)
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font_name
    # Set complex script font for Hebrew
    rPr = run._r.get_or_add_rPr()
    cs_font = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="{font_name}"/>')
    rPr.append(cs_font)
    set_rtl_run(run)
    return run


def add_english_run(paragraph, text, bold=False, size=Pt(11), color=DARK_GRAY, font_name='Courier New'):
    """Add an English/LTR text run."""
    run = paragraph.add_run(text)
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = font_name
    return run


def set_section_rtl(section):
    """Set the section to RTL document direction."""
    sectPr = section._sectPr
    bidi = parse_xml(f'<w:bidi {nsdecls("w")} val="1"/>')
    sectPr.append(bidi)


def add_heading_rtl(doc, text, level=1, color=DARK_BLUE):
    """Add a RTL heading."""
    heading = doc.add_heading(level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(heading)
    run = heading.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Arial'
    rPr = run._r.get_or_add_rPr()
    cs_font = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="Arial"/>')
    rPr.append(cs_font)
    set_rtl_run(run)
    return heading


def add_paragraph_rtl(doc, text='', bold=False, italic=False, size=Pt(11), color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(6)):
    """Add a RTL paragraph with text."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    set_rtl_paragraph(p)
    if text:
        add_hebrew_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_screenshot(doc, filename, caption):
    """Add a screenshot image with caption."""
    img_path = SCREENSHOTS_DIR / filename
    if not img_path.exists():
        print(f"  WARNING: Screenshot not found: {filename}")
        return

    # Re-save through Pillow to fix ffmpeg JPEG headers
    img = PILImage.open(img_path)
    buf = io.BytesIO()
    img.save(buf, format='JPEG', quality=85)
    buf.seek(0)

    # Image paragraph (centered)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(buf, width=Inches(5.5))

    # Caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    set_rtl_paragraph(p_cap)
    add_hebrew_run(p_cap, caption, italic=True, size=Pt(9), color=MEDIUM_GRAY)


def add_table_rtl(doc, headers, rows, col_widths=None):
    """Add a RTL table with headers and data rows."""
    # Reverse headers and rows for RTL display (rightmost column first)
    headers_rtl = list(reversed(headers))
    rows_rtl = [list(reversed(row)) for row in rows]

    table = doc.add_table(rows=1 + len(rows_rtl), cols=len(headers_rtl))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set RTL on table
    tblPr = table._tbl.tblPr
    bidi_elem = parse_xml(f'<w:bidiVisual {nsdecls("w")} val="1"/>')
    tblPr.append(bidi_elem)

    # Header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers_rtl):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl_paragraph(p)
        add_hebrew_run(p, header_text, bold=True, size=Pt(9), color=WHITE)
        # Dark blue background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0A3D62" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for row_idx, row_data in enumerate(rows_rtl):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl_paragraph(p)
            # Check if text should be bold
            if cell_text.startswith('**') and cell_text.endswith('**'):
                add_hebrew_run(p, cell_text[2:-2], bold=True, size=Pt(9))
            else:
                add_hebrew_run(p, cell_text, size=Pt(9))
            # Alternating row colors
            if row_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Add space after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def add_info_box(doc, title, content_lines, border_color='0077B6', bg_color='E8F4F8'):
    """Add a styled info box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(p)

    # Add right border and shading via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:right w:val="single" w:sz="24" w:space="8" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
    pPr.append(shading)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    if title:
        add_hebrew_run(p, title, bold=True, size=Pt(10), color=BRAND_BLUE)
        run = p.add_run('\n')

    for i, line in enumerate(content_lines):
        add_hebrew_run(p, line, size=Pt(10))
        if i < len(content_lines) - 1:
            p.add_run('\n')


def add_warning_box(doc, title, content_lines):
    """Add a warning-style box."""
    add_info_box(doc, title, content_lines, border_color='F39C12', bg_color='FFF3CD')


def add_phase_box(doc, title, items):
    """Add a phase/milestone box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(p)

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:right w:val="single" w:sz="24" w:space="8" w:color="27AE60"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F7F0" w:val="clear"/>')
    pPr.append(shading)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    add_hebrew_run(p, title, bold=True, size=Pt(11))
    for item in items:
        p.add_run('\n')
        add_hebrew_run(p, f'• {item}', size=Pt(10))


def add_bullet_list(doc, items):
    """Add a bulleted list."""
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl_paragraph(p)
        # Check if item has bold prefix
        if '**' in item:
            parts = item.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Bold part
                    add_hebrew_run(p, part, bold=True, size=Pt(10))
                else:
                    if part:
                        add_hebrew_run(p, part, size=Pt(10))
        else:
            add_hebrew_run(p, item, size=Pt(10))


def add_numbered_list(doc, items):
    """Add a numbered list."""
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl_paragraph(p)
        if '**' in item:
            parts = item.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    add_hebrew_run(p, part, bold=True, size=Pt(10))
                else:
                    if part:
                        add_hebrew_run(p, part, size=Pt(10))
        else:
            add_hebrew_run(p, item, size=Pt(10))


def add_diagram(doc, text):
    """Add a monospace diagram block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="DDDDDD"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="DDDDDD"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="4" w:color="DDDDDD"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="4" w:color="DDDDDD"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F6FA" w:val="clear"/>')
    pPr.append(shading)

    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)


def add_decision_box(doc, title, content_lines):
    """Add a decision-style box (purple)."""
    add_info_box(doc, title, content_lines, border_color='7B2D8E', bg_color='F3E8F9')


def add_critical_box(doc, title, content_lines):
    """Add a critical-style box (red)."""
    add_info_box(doc, title, content_lines, border_color='C0392B', bg_color='FDEDEC')


def add_quote_box(doc, speaker, text):
    """Add a quote box with speaker name bold, rest italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(p)

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:right w:val="single" w:sz="24" w:space="8" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'  )
    pPr.append(shading)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    add_hebrew_run(p, speaker, bold=True, size=Pt(10), color=MEDIUM_GRAY)
    add_hebrew_run(p, f' {text}', italic=True, size=Pt(10))


def add_phase_box_yellow(doc, title, items):
    """Add a yellow phase box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(p)

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:right w:val="single" w:sz="24" w:space="8" w:color="F39C12"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF3CD" w:val="clear"/>'  )
    pPr.append(shading)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    add_hebrew_run(p, title, bold=True, size=Pt(11))
    for item in items:
        p.add_run('\n')
        add_hebrew_run(p, f'• {item}', size=Pt(10))


def add_phase_box_orange(doc, title, items):
    """Add an orange phase box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(p)

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:right w:val="single" w:sz="24" w:space="8" w:color="E67E22"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FDEBD0" w:val="clear"/>'  )
    pPr.append(shading)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    add_hebrew_run(p, title, bold=True, size=Pt(11))
    for item in items:
        p.add_run('\n')
        add_hebrew_run(p, f'• {item}', size=Pt(10))


def add_phase_box_red(doc, title, items):
    """Add a red phase box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl_paragraph(p)

    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:right w:val="single" w:sz="24" w:space="8" w:color="C0392B"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FDEDEC" w:val="clear"/>'  )
    pPr.append(shading)

    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    add_hebrew_run(p, title, bold=True, size=Pt(11))
    for item in items:
        p.add_run('\n')
        add_hebrew_run(p, f'• {item}', size=Pt(10))



def build_document():
    """Build the full Hebrew RTL Word document - v2.0 (12-section, meeting-driven)."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = DARK_GRAY
    rPr = style.element.get_or_add_rPr()
    cs_font = parse_xml(f'<w:rFonts {nsdecls("w")} w:cs="Arial"/>')
    rPr.append(cs_font)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    set_section_rtl(section)

    # ==================== COVER PAGE ====================
    for _ in range(6):
        doc.add_paragraph()
    add_paragraph_rtl(doc, 'TALIAZ HEALTH • HealthyMind', bold=True, size=Pt(14), color=BRAND_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph_rtl(doc, 'מסמך מפרט טכני', bold=True, size=Pt(28), color=DARK_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_rtl(doc, 'אינטגרציית מערכת קליקס (Clicks EHR)', bold=True, size=Pt(22), color=DARK_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph_rtl(doc, 'תליאז — מכבי שירותי בריאות', size=Pt(16), color=BRAND_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="1" w:color="0077B6"/></w:pBdr>')
    pPr.append(pBdr)
    doc.add_paragraph()
    meta_items = [
        ('הוכן על ידי:', 'ד"ר ראמי ח\'ורי, CMIO — תליאז הלת\' / מרפאת HealthyMind'),
        ('תאריך:', '22 במרץ 2026'),
        ('גרסה:', '2.0'),
        ('נמענים:', 'שלומי (מוביל טכני, מכבי), יהודה (ראש אגף מערכות מידע, מכבי), צוות האדריכלות של מכבי'),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl_paragraph(p)
        add_hebrew_run(p, label, bold=True, size=Pt(11), color=RGBColor(0x33, 0x33, 0x33))
        add_hebrew_run(p, f' {value}', size=Pt(11), color=MEDIUM_GRAY)
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_paragraph(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="12" w:space="4" w:color="C0392B"/><w:bottom w:val="single" w:sz="12" w:space="4" w:color="C0392B"/><w:left w:val="single" w:sz="12" w:space="4" w:color="C0392B"/><w:right w:val="single" w:sz="12" w:space="4" w:color="C0392B"/></w:pBdr>')
    pPr.append(pBdr)
    add_hebrew_run(p, 'סודי — לצוות הטכני של מכבי בלבד', bold=True, size=Pt(12), color=RED)
    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_heading_rtl(doc, 'תוכן עניינים', level=1)
    toc_items = [
        (1, '1. תקציר מנהלים'),
        (1, '2. רקע — ישיבת התיאום והחלטות'),
        (2, '2.1 משתתפים'), (2, '2.2 החלטות מרכזיות'), (2, '2.3 פריטי פעולה'),
        (1, '3. סקירת ארכיטקטורת המערכת'),
        (1, '4. מנגנוני אינטגרציה קיימים'),
        (2, '4.1 הכספת — ניתוח מעמיק'), (2, '4.2 פרויקט בת קול'), (2, '4.3 WebAge Wrapper'),
        (1, '5. תיעוד מלא של זרימת העבודה — מסך אחר מסך'),
        (2, '5.1-5.9 תת-סעיפים'),
        (1, '6. סיכום זרימת נתונים'),
        (1, '7. אבטחת מידע וגישה מרחוק'),
        (1, '8. אינטגרציית iFeel וסיכומים פסיכיאטריים קודמים'),
        (1, '9. חוק ניוד המידע ו-FHIR'),
        (1, '10. מחיצת בריאות הנפש'),
        (1, '11. מקרי בדיקה ושרשרת אימיילים'),
        (1, '12. שלבי אינטגרציה ופריטי פעולה'),
    ]
    for level, text in toc_items:
        indent = '    ' if level == 2 else ''
        size = Pt(10.5) if level == 2 else Pt(11)
        p = add_paragraph_rtl(doc, f'{indent}{text}', size=size, space_after=Pt(4))
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="dotted" w:sz="4" w:space="1" w:color="CCCCCC"/></w:pBdr>')
        pPr.append(pBdr)
    doc.add_page_break()


    # ==================== SECTION 1: EXECUTIVE SUMMARY ====================
    add_heading_rtl(doc, '1. תקציר מנהלים', level=1)
    add_paragraph_rtl(doc, 'מסמך זה מתאר את המפרט הטכני המלא לאינטגרציה בין הפלטפורמה הקלינית של תליאז (Predictix/Xano) לבין מערכת קליקס EHR של מכבי שירותי בריאות. המסמך מבוסס על ישיבת תיאום שנערכה ב-19 במרץ 2026, על הקלטת מסך מלאה של זרימת העבודה בקליקס (נובמבר 2024), ועל שרשרת אימייל מפורטת שתיעדה את ניסיונות האינטגרציה הקודמים.')
    add_heading_rtl(doc, 'מטרה', level=3)
    add_paragraph_rtl(doc, 'להחליף את הזנת הנתונים הידנית בקליקס בקריאות API אוטומטיות מהפלטפורמה הקלינית של תליאז, ובכך לאפשר:')
    add_bullet_list(doc, [
        '**העלאת סיכום אוטומטית** (תשובת יועץ) מתליאז לקליקס',
        '**הוצאת מרשמים** באמצעות API במקום הזנה ידנית',
        '**סנכרון נתונים דו-כיווני** — שליפת נתוני דמוגרפיה, תרופות ותוצאות מעבדה לתוך תליאז',
    ])
    add_info_box(doc, 'גרסאות מערכת שנצפו:', [
        'W-Clicks Dot.NET | Version 2023.12.109.23696',
        'Application version 2024.09.101',
        'גישה דרך Citrix StoreFront בכתובת: d.mac.org.il/Citrix/ExtWeb/',
    ])
    add_decision_box(doc, 'החלטה מרכזית מהישיבה:', ['שלומי (מוביל טכני, מכבי) אישר שניתן לספק לתליאז גישת API לכתיבה וקריאה מקליקס, בכפוף לבחינה טכנית של צוות האדריכלות. המסמך הנוכחי נכתב כדי לאפשר את הבחינה הזו.'])

    # ==================== SECTION 2: MEETING BACKGROUND ====================
    doc.add_page_break()
    add_heading_rtl(doc, '2. רקע — ישיבת התיאום והחלטות', level=1)
    add_paragraph_rtl(doc, 'ב-19 במרץ 2026 נערכה ישיבת תיאום בין תליאז למכבי לדיון באפשרויות האינטגרציה. הישיבה כללה סקירה של המצב הנוכחי, ניסיונות העבר (כספת), ומיפוי דרישות עתידיות.')
    add_heading_rtl(doc, '2.1 משתתפים', level=2, color=BRAND_BLUE)
    add_table_rtl(doc, ['שם', 'תפקיד', 'ארגון'], [
        ['שלומי', 'מוביל טכני, מערכות מידע', 'מכבי'],
        ['יהודה', 'ראש אגף מערכות מידע', 'מכבי'],
        ['ד"ר ראמי ח׳ורי', 'CMIO', 'תליאז / HealthyMind'],
        ['דקל תליאז', 'מנכ"ל', 'תליאז'],
        ['שחר שגב', 'מנהלת תפעול קליני', 'תליאז'],
        ['ענת מירון', 'מנהלת בריאות הנפש', 'מכבי'],
    ])
    add_heading_rtl(doc, '2.2 החלטות מרכזיות', level=2, color=BRAND_BLUE)
    add_decision_box(doc, 'החלטה 1 — גישת API:', ['שלומי אישר עקרונית שניתן לספק לתליאז API לקליקס. נדרשת בחינה טכנית של צוות האדריכלות על בסיס המסמך הנוכחי.'])
    add_decision_box(doc, 'החלטה 2 — הכספת:', ['מנגנון הכספת הקיים אינו מספק — סיכומים מועלים אך אינם מנותבים לרופא המשפחה. יש לבחון מחדש את התצורה או לעבור ל-API ישיר.'])
    add_decision_box(doc, 'החלטה 3 — פרויקט בת קול כמודל:', ['ל"בת קול" כבר יש API להעלאת סיכומים לקליקס. תשתית זו יכולה לשמש כבסיס לאינטגרציה של תליאז.'])
    add_decision_box(doc, 'החלטה 4 — סביבת בדיקות:', ['שלומי הציע שתליאז תקבל גישה לסביבת בדיקות (TEST environment) של קליקס לפיתוח ובדיקת האינטגרציה.'])
    add_decision_box(doc, 'החלטה 5 — מסמך טכני:', ['תליאז תכין מסמך מפרט טכני מפורט (המסמך הנוכחי) לצוות האדריכלות של מכבי, כבסיס לתכנון ה-API.'])
    add_warning_box(doc, '⚠ אזהרת ענת מירון:', ['ענת מירון, מנהלת בריאות הנפש במכבי, הדגישה שכל שינוי בזרימת העבודה חייב לקבל את אישורה ואת אישור הנהלת בריאות הנפש. היא ביקשה להיות מעורבת בכל שלבי התכנון.'])
    add_heading_rtl(doc, '2.3 פריטי פעולה', level=2, color=BRAND_BLUE)
    add_quote_box(doc, 'שלומי (מכבי) — מהישיבה:', '"אני אצוות מישהו, אני אביס לכם ארכיטקט, שיישב איתכם [...] זה לא צריך לתמחר את זה, ולא צריך לתעדף את זה, יהודה, זה עליי."')
    add_quote_box(doc, 'שלומי (מכבי) — מהישיבה:', '"אני מציע, אני אשים ארכיטקט, יחד עם טכנולוגיות רפואיות, מישהו מטכנולוגיות רפואיות, יחד עם ארכיטקט, ישבו רגע איתכם להבין על מה מדובר [...] הם יבואו אלינו עם סיכום של מה זה אומר, מה המשמעות, מה קיים היום, ומה אפשר לעשות."')
    add_quote_box(doc, 'שלומי (מכבי) — מהישיבה:', '"שימו הכל, שימו את הכל על מסמך [...] כל החלומות, תשימו מה שאתם רוצים, הכל בפנים, ואז נבין מה קיים [...] נשים ארכיטקט חינם, בסדר, כרגע לא עולה כסף, ארכיטקט, נשים מנתח מערכות, מאפיין בצד של טכנולוגיות רפואיות, שמכיר את הקליקס טוב."')
    add_info_box(doc, 'סיכום:', ['שלומי התחייב לצוות ארכיטקט + מנתח מערכות מטכנולוגיות רפואיות (שני אנשים), ללא עלות בשלב זה. תליאז מתבקשת להכין מסמך מפורט עם "כל החלומות" — כל הדרישות, כל המסכים, כל השדות. ישיבת קבלת החלטות לפני פסח.'])
    add_table_rtl(doc, ['#', 'פריט', 'אחראי', 'לוח זמנים'], [
        ['1', 'הכנת מסמך מפרט טכני מלא — "כל החלומות"', 'ד"ר ראמי ח׳ורי + שחר שגב (תליאז)', 'מיידי'],
        ['2', 'שיבוץ ארכיטקט + מנתח מערכות (טכנולוגיות רפואיות)', 'שלומי (מכבי)', 'מיידי — ללא עלות'],
        ['3', 'פגישת הערכה טכנית (ארכיטקט + תליאז)', 'שלומי + שחר שגב', 'שבוע-שבועיים'],
        ['4', 'ישיבת קבלת החלטות עם יהודה', 'שלומי + יהודה + תליאז', 'לפני פסח'],
        ['5', 'בדיקת פתרון Remote Desktop מחו״ל מול אבטחת מידע', 'שלומי (מכבי)', 'לאחר קבלת סיכום מאת שחר'],
        ['6', 'שליחת סיכום פגישה לשלומי', 'שחר שגב (תליאז)', 'מיידי'],
    ])


    # ==================== SECTION 3: ARCHITECTURE ====================
    doc.add_page_break()
    add_heading_rtl(doc, '3. סקירת ארכיטקטורת המערכת', level=1)
    add_heading_rtl(doc, 'מצב נוכחי', level=3)
    add_diagram(doc, "Current State\n\n  Psychiatrist ---Citrix VPN---> Clicks EHR\n  (Local PC)                     (Maccabi servers)\n       | Copy-Paste\n       v\n  Predictix (Xano/Web)\n  Taliaz Clinical Platform")
    add_heading_rtl(doc, 'מצב יעד', level=3)
    add_diagram(doc, "Target State\n\n  Predictix  <---API Layer--->  Clicks EHR\n  (Xano/Web)    (REST/SOAP)     (Maccabi servers)\n       |                              |\n       +------ Automated Flow --------+\n         - Upload summaries\n         - Sync prescriptions\n         - Pull patient data")
    add_quote_box(doc, 'שלומי (מכבי):', '"אם תכינו מסמך מפורט עם כל השדות והזרימות, אני אעביר את זה לצוות האדריכלות שלנו..."')

    # ==================== SECTION 4: EXISTING MECHANISMS ====================
    doc.add_page_break()
    add_heading_rtl(doc, '4. מנגנוני אינטגרציה קיימים', level=1)
    add_paragraph_rtl(doc, 'לפני תכנון אינטגרציה חדשה, יש לסקור את המנגנונים הקיימים שכבר פועלים (או שנוסו) בין ספקים חיצוניים לקליקס:')
    add_heading_rtl(doc, '4.1 הכספת — ניתוח מעמיק', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, 'הכספת היא מנגנון קיים במכבי המאפשר לספקים חיצוניים להעלות מסמכים לתיק המטופל בקליקס. תליאז ניסתה להשתמש בכספת להעלאת סיכומי ייעוץ — אך נתקלה בבעיות חמורות.')
    add_critical_box(doc, '3 בעיות קריטיות שזוהו בכספת:', [
        '1. ניתוב שגוי: סיכומים הועלו בהצלחה לכספת אך לא הגיעו לדואר הרפואי של רופא המשפחה.',
        '2. היעדר תשובת יועץ: הכספת מעלה מסמך כקובץ מצורף, אך אינה יוצרת "תשובת יועץ" רשמית בקליקס.',
        '3. דחייה על ידי ענת מירון: מנהלת בריאות הנפש במכבי דחתה את מנגנון הכספת כפתרון קבוע.',
    ])
    add_heading_rtl(doc, 'השוואת כספת מול API ישיר', level=3)
    add_table_rtl(doc, ['קריטריון', 'כספת', 'API ישיר לקליקס'], [
        ['העלאת סיכום', 'קובץ מצורף בלבד', 'תשובת יועץ רשמית'],
        ['התראה לרופא משפחה', 'לא', 'כן'],
        ['קודי אבחנה', 'לא', 'כן'],
        ['מרשמים', 'לא', 'אפשרי'],
        ['שליפת נתוני מטופל', 'לא', 'כן'],
        ['מורכבות הקמה', 'נמוכה', 'גבוהה'],
        ['תחזוקה שוטפת', 'נמוכה', 'בינונית'],
    ])
    add_heading_rtl(doc, '4.2 פרויקט בת קול', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, 'פרויקט בת קול של מכבי הוא מערכת קיימת שכבר משתמשת ב-API להעלאת סיכומים לקליקס. בישיבה, שלומי הזכיר את בת קול כמודל אפשרי לאינטגרציה של תליאז.')
    add_quote_box(doc, 'שלומי (מכבי):', '"בת קול כבר עושים את זה — יש להם API שכותב לקליקס. אפשר לבדוק אם אפשר לעשות את אותו דבר עבורכם."')
    add_info_box(doc, 'משמעות:', ['קיומו של API עובד בבת קול מוכיח שהדבר אפשרי טכנית. השאלה היא האם ניתן להרחיב את אותה תשתית לשימוש תליאז, או שנדרש API נפרד.'])
    add_heading_rtl(doc, '4.3 WebAge Wrapper', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, 'שלומי הזכיר בישיבה את WebAge — שכבת wrapper שמכבי משתמשת בה כממשק בין מערכות חיצוניות לקליקס. ייתכן שזו שכבת ה-middleware שדרכה ניתן יהיה לבצע את האינטגרציה.')
    add_quote_box(doc, 'שלומי (מכבי):', '"יש לנו WebAge שיושב מעל קליקס ומאפשר תקשורת. צריך לבדוק מה אפשר לחשוף דרכו."')


    # ==================== SECTION 5: WORKFLOW (was section 3) ====================
    doc.add_page_break()
    add_heading_rtl(doc, '5. תיעוד מלא של זרימת העבודה — מסך אחר מסך', level=1)

    # 5.1 Access
    add_heading_rtl(doc, '5.1 גישה והזדהות', level=2, color=BRAND_BLUE)
    add_heading_rtl(doc, 'שלב 1: כניסה דרך Citrix StoreFront', level=3)
    add_screenshot(doc, '01_citrix_storefront_app_launcher.jpg', 'צילום מסך 1: פורטל Citrix StoreFront — השקת אפליקציית קליקס')
    add_paragraph_rtl(doc, 'הפסיכיאטר ניגש לקליקס דרך פורטל Citrix StoreFront של מכבי בכתובת d.mac.org.il/Citrix/ExtWeb/. הפורטל מציג 19 אפליקציות זמינות, כולל:')
    add_bullet_list(doc, ['**ניהול תיק רפואי-2017** — זו אפליקציית קליקס', 'Citrix Excel, Word, PowerPoint, Outlook', 'Bomgar (תמיכה מרחוק)', 'פורטלים שונים של מכבי'])
    add_info_box(doc, 'הערת אינטגרציה:', ['שכבת ה-Citrix מוסיפה מורכבות — כל API יצטרך לעקוף את סשן ה-Citrix או לפעול בצד השרת ברשת הפנימית של מכבי.'])

    # 5.2 Doctor Portal
    doc.add_page_break()
    add_heading_rtl(doc, '5.2 פורטל רופא ובחירת מטופל', level=2, color=BRAND_BLUE)
    add_heading_rtl(doc, 'שלב 2: פורטל רופא בקליקס', level=3)
    add_screenshot(doc, '02_clicks_doctor_portal_patient_list.jpg', 'צילום מסך 2: פורטל רופא עם רשימת מטופלים ומערכת חיפוש')
    add_paragraph_rtl(doc, 'לאחר ההשקה, הרופא נכנס לפורטל הרופא. אזורים עיקריים:')
    add_table_rtl(doc, ['שדה', 'שם עברי', 'סוג', 'הערות'], [
        ['ת.ז. מטופל', 'תעודת זהות', 'מספר בן 9 ספרות', 'מפתח ראשי'],
        ['טלפון', 'טלפון', 'מחרוזת', ''],
        ['שם פרטי', 'שם פרטי', 'מחרוזת', ''],
        ['שם משפחה', 'שם משפחה', 'מחרוזת', ''],
    ])
    add_heading_rtl(doc, 'שלב 3: תצוגה מקבילה — קליקס + Predictix של תליאז', level=3)
    add_screenshot(doc, '03_clicks_and_predictix_side_by_side.jpg', 'צילום מסך 3: עבודה במקביל עם קליקס (למעלה) ופלטפורמת Predictix של תליאז (למטה)')
    add_info_box(doc, 'הזדמנות אינטגרציה:', ['ביטול ה-copy-paste בין שתי המערכות על ידי push/pull אוטומטי של נתונים מ-Predictix לקליקס באמצעות API.'])

    # 5.3 Patient File
    doc.add_page_break()
    add_heading_rtl(doc, '5.3 תיק מטופל ונתונים קליניים', level=2, color=BRAND_BLUE)
    add_heading_rtl(doc, 'שלב 4: פרטים אישיים של המטופל', level=3)
    add_screenshot(doc, '04_patient_file_personal_details.jpg', 'צילום מסך 4: תיק מטופל — פרטים אישיים, בעיות ידועות, תרופות קבועות וגורמי סיכון')
    add_paragraph_rtl(doc, 'לאחר בחירת מטופל, נפתח התיק המלא בקליקס:')
    add_heading_rtl(doc, 'שדות לגישת API — פעולות קריאה (READ):', level=4)
    add_table_rtl(doc, ['שדה', 'שם עברי', 'עדיפות API'], [
        ['דמוגרפיה', 'פרטים אישיים', '**גבוהה**'],
        ['אבחנות ידועות', 'בעיות ידועות', '**גבוהה**'],
        ['תרופות נוכחיות', 'תרופות קבועות', '**גבוהה**'],
        ['אלרגיות/רגישויות', 'רגישויות', '**גבוהה**'],
        ['תוצאות מעבדה', 'מעבדות', 'בינונית'],
        ['ייעוצים קודמים', 'ייעוצים', 'בינונית'],
        ['גורמי סיכון', 'גורמי סיכון', 'נמוכה'],
    ])


    # 5.4 Visit Registration
    doc.add_page_break()
    add_heading_rtl(doc, '5.4 רישום ביקורים', level=2, color=BRAND_BLUE)
    add_heading_rtl(doc, 'שלב 5: פתיחת ביקור', level=3)
    add_screenshot(doc, '05_visit_registration_screen.jpg', 'צילום מסך 5: מסך רישום ביקורים — פתיחת ייעוץ פסיכיאטרי')
    add_heading_rtl(doc, 'שלב 6: הזנת ממצאים קליניים', level=3)
    add_screenshot(doc, '06_clinical_findings_form.jpg', 'צילום מסך 6: טופס ממצאים קליניים עם קטגוריות בדיקה')
    add_heading_rtl(doc, 'שדות לגישת API — פעולות כתיבה (WRITE):', level=4)
    add_table_rtl(doc, ['שדה', 'שם עברי', 'סוג', 'עדיפות'], [
        ['תאריך ושעת ביקור', 'תאריך ושעה', 'DateTime', '**גבוהה**'],
        ['מזהה רופא', 'מזהה רופא', 'מחרוזת', '**גבוהה**'],
        ['סוג ביקור', 'סוג ביקור', 'אינטייק/מעקב', '**גבוהה**'],
        ['ממצאים קליניים', 'ממצאי הבדיקה', 'טקסט חופשי', '**גבוהה**'],
        ['הערכת מסוכנות', 'מסוכנות', 'תיבות סימון', '**גבוהה**'],
        ['סיבת הפניה', 'סיבת הפניה', 'מחרוזת', 'בינונית'],
    ])

    # 5.5 Diagnosis
    doc.add_page_break()
    add_heading_rtl(doc, '5.5 הזנת אבחנה', level=2, color=BRAND_BLUE)
    add_heading_rtl(doc, 'שלב 7: חיפוש ובחירת אבחנה (ICD)', level=3)
    add_screenshot(doc, '07_diagnosis_icd_entry.jpg', 'צילום מסך 7: השלמה אוטומטית של קודי ICD — חיפוש "POST TR"')
    add_info_box(doc, 'הערה:', ['קליקס משתמש בקודים פנימיים קצרים הממופים ל-ICD. כל API יצטרך להשתמש בקודים הפנימיים הללו, לא בקודי ICD-10 סטנדרטיים.'])
    add_heading_rtl(doc, 'שלב 8: בחירת סוג טיפול (קודים פסיכיאטריים)', level=3)
    add_screenshot(doc, '08_treatment_type_psychiatric_codes.jpg', 'צילום מסך 8: קודי שירות פסיכיאטרי במערכת קליקס')

    # 5.6 Consultant Response
    doc.add_page_break()
    add_heading_rtl(doc, '5.6 תשובת יועץ (מסמך הסיכום)', level=2, color=BRAND_BLUE)
    add_warning_box(doc, '⚠ יעד מס\' 1 לאוטומציה:', ['תשובת היועץ היא המנגנון העיקרי לתקשורת הממצאים הפסיכיאטריים בחזרה לרופא המשפחה המפנה. זהו היעד בעל ה-ROI הגבוה ביותר לאוטומציה באמצעות API.'])
    add_heading_rtl(doc, 'שלב 9: בחירת הרופא המפנה', level=3)
    add_screenshot(doc, '09_consultant_response_doctor_selection.jpg', 'צילום מסך 9: דיאלוג תשובת יועץ — בחירת רופא מטפל להעברת התשובה')
    add_heading_rtl(doc, 'שלב 10: תוכן מכתב תשובת היועץ', level=3)
    add_screenshot(doc, '10_consultant_response_letter_content.jpg', 'צילום מסך 10: מכתב תשובת יועץ מלא עם פרטי מטופל וטקסט קליני')
    add_heading_rtl(doc, 'שדות לגישת API — כתיבת סיכום (עדיפות קריטית):', level=4)
    add_table_rtl(doc, ['שדה', 'סוג', 'הערות'], [
        ['תאריך בדיקה', 'תאריך', 'בדיקה שנעשתה ב...'],
        ['סוג ביקור', 'מחרוזת', 'אינטייק / מעקב'],
        ['נרטיב קליני', 'RTF/HTML', 'טקסט הערכה מלא'],
        ['אבחנה', 'ICD-10 + טקסט', 'F43.1 וכו\''],
        ['המלצות טיפול', 'טקסט', 'תרופות מומלצות'],
        ['מזהה רופא מפנה', 'מחרוזת', 'יעד לתשובה'],
    ])
    add_heading_rtl(doc, 'שלב 11: אבחנה והמלצות', level=3)
    add_screenshot(doc, '11_summary_diagnosis_recommendations.jpg', 'צילום מסך 11: אישור אבחנה, תרופות מומלצות ואפשרויות חתימה')

    # 5.7 Visit Completion
    doc.add_page_break()
    add_heading_rtl(doc, '5.7 סיום ביקור וחתימה', level=2, color=BRAND_BLUE)
    add_heading_rtl(doc, 'שלב 12: קוד טיפול ותזמון מעקב', level=3)
    add_screenshot(doc, '12_visit_completion_treatment_code.jpg', 'צילום מסך 12: קוד טיפול PSYCHIATRIC INTAKE - ADULT (9864), תזמון מעקב')
    add_heading_rtl(doc, 'שלב 13: שליחת טפסים והנחיות למטופל', level=3)
    add_screenshot(doc, '13_form_submission_patient_instructions.jpg', 'צילום מסך 13: דיאלוג שליחת טפסים — אישור ותשובת יועץ')
    add_heading_rtl(doc, 'שלב 14: חתימה דיגיטלית', level=3)
    add_screenshot(doc, '14_digital_signature_login.jpg', 'צילום מסך 14: דיאלוג הזדהות לחתימה דיגיטלית')
    add_paragraph_rtl(doc, 'החתימה הדיגיטלית דורשת שם משתמש וסיסמה. זוהי שכבת אבטחה קריטית — כל אינטגרציית API תצטרך לטפל בהזדהות וחתימה באופן תכנותי.')
    add_heading_rtl(doc, 'שלב 15: רשומת ביקור חתומה', level=3)
    add_screenshot(doc, '15_signed_visit_record.jpg', 'צילום מסך 15: רשומת ביקור חתומה עם חותמת זמן ותוכן קליני מלא')

    # 5.8 Prescriptions
    doc.add_page_break()
    add_heading_rtl(doc, '5.8 ניהול מרשמים', level=2, color=BRAND_BLUE)
    add_heading_rtl(doc, 'שלב 16: תצוגת תרופות קבועות', level=3)
    add_screenshot(doc, '16_chronic_medications_view.jpg', 'צילום מסך 16: טבלת תרופות קבועות — סטטוס מרשמים נוכחי')
    add_heading_rtl(doc, 'שלב 17: עדכון מינון תרופה', level=3)
    add_screenshot(doc, '17_medication_dosage_update.jpg', 'צילום מסך 17: דיאלוג עדכון מינון — MIRO 30MG')
    add_heading_rtl(doc, 'שלב 18: חיפוש תרופה חדשה', level=3)
    add_screenshot(doc, '18_new_prescription_drug_search.jpg', 'צילום מסך 18: חיפוש תרופה (LORIVAN) עם התרעות ומידע על מטופל')
    add_heading_rtl(doc, 'שלב 19: טופס מרשם — מינון והוראות', level=3)
    add_screenshot(doc, '19_prescription_form_dosage_instructions.jpg', 'צילום מסך 19: טופס מרשם — LORIVAN TAB 1MG X 20 עם הוראות מינון')
    add_heading_rtl(doc, 'שלב 20: אישור שמירת מרשם', level=3)
    add_screenshot(doc, '20_prescription_save_confirmation.jpg', 'צילום מסך 20: דיאלוג אישור שמירת מרשם')
    add_heading_rtl(doc, 'שלב 21: הפקת מרשמים', level=3)
    add_screenshot(doc, '21_prescription_generation_screen.jpg', 'צילום מסך 21: מסך הפקת מרשמים — בחירת תרופות, משך ואפשרויות שליחה')

    # 5.9 External System
    doc.add_page_break()
    add_heading_rtl(doc, '5.9 מערכת חיצונית — מסמך סיכום', level=2, color=BRAND_BLUE)
    add_screenshot(doc, '24_google_docs_summary_document.jpg', 'צילום מסך 22: מסמך סיכום ב-Google Docs — מוכן ל-copy-paste לקליקס')
    add_warning_box(doc, '⚠ הערת אינטגרציה:', ['זהו בדיוק התוכן שצריך להידחף אוטומטית מתליאז לקליקס באמצעות API, ובכך לבטל את זרימת ה-copy-paste בין Google Docs לתשובת יועץ בקליקס.'])


    # ==================== SECTION 6: DATA FLOW SUMMARY ====================
    doc.add_page_break()
    add_heading_rtl(doc, '6. \u05e1\u05d9\u05db\u05d5\u05dd \u05d6\u05e8\u05d9\u05de\u05ea \u05e0\u05ea\u05d5\u05e0\u05d9\u05dd \u2014 \u05e9\u05d3\u05d5\u05ea \u05d4\u05d3\u05d5\u05e8\u05e9\u05d9\u05dd \u05d2\u05d9\u05e9\u05ea API', level=1)
    add_heading_rtl(doc, '6.1 \u05e4\u05e2\u05d5\u05dc\u05d5\u05ea \u05e7\u05e8\u05d9\u05d0\u05d4 (\u05e7\u05dc\u05d9\u05e7\u05e1 \u2190 \u05ea\u05dc\u05d9\u05d0\u05d6)', level=2, color=BRAND_BLUE)
    add_table_rtl(doc, ['\u05e2\u05d3\u05d9\u05e4\u05d5\u05ea', '\u05e0\u05ea\u05d5\u05df', '\u05de\u05e1\u05da \u05d1\u05e7\u05dc\u05d9\u05e7\u05e1', '\u05e9\u05d9\u05d8\u05d4 \u05e0\u05d5\u05db\u05d7\u05d9\u05ea'], [
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05d3\u05de\u05d5\u05d2\u05e8\u05e4\u05d9\u05d4', '\u05e4\u05e8\u05d8\u05d9\u05dd \u05d0\u05d9\u05e9\u05d9\u05d9\u05dd', '\u05d7\u05d9\u05e4\u05d5\u05e9 \u05d9\u05d3\u05e0\u05d9'],
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05d0\u05d1\u05d7\u05e0\u05d5\u05ea \u05d9\u05d3\u05d5\u05e2\u05d5\u05ea', '\u05d1\u05e2\u05d9\u05d5\u05ea \u05d9\u05d3\u05d5\u05e2\u05d5\u05ea', '\u05e7\u05e8\u05d9\u05d0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05ea\u05e8\u05d5\u05e4\u05d5\u05ea \u05e0\u05d5\u05db\u05d7\u05d9\u05d5\u05ea', '\u05ea\u05e8\u05d5\u05e4\u05d5\u05ea \u05e7\u05d1\u05d5\u05e2\u05d5\u05ea', '\u05e7\u05e8\u05d9\u05d0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05d0\u05dc\u05e8\u05d2\u05d9\u05d5\u05ea', '\u05e8\u05d2\u05d9\u05e9\u05d5\u05d9\u05d5\u05ea', '\u05e7\u05e8\u05d9\u05d0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['\u05d1\u05d9\u05e0\u05d5\u05e0\u05d9\u05ea', '\u05ea\u05d5\u05e6\u05d0\u05d5\u05ea \u05de\u05e2\u05d1\u05d3\u05d4', '\u05de\u05e2\u05d1\u05d3\u05d5\u05ea', '\u05e7\u05e8\u05d9\u05d0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['\u05d1\u05d9\u05e0\u05d5\u05e0\u05d9\u05ea', '\u05d9\u05d9\u05e2\u05d5\u05e6\u05d9\u05dd \u05e7\u05d5\u05d3\u05de\u05d9\u05dd', '\u05d9\u05d9\u05e2\u05d5\u05e6\u05d9\u05dd', '\u05e7\u05e8\u05d9\u05d0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['\u05e0\u05de\u05d5\u05db\u05d4', '\u05d2\u05d5\u05e8\u05de\u05d9 \u05e1\u05d9\u05db\u05d5\u05df', '\u05d2\u05d5\u05e8\u05de\u05d9 \u05e1\u05d9\u05db\u05d5\u05df', '\u05e7\u05e8\u05d9\u05d0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
    ])
    add_heading_rtl(doc, '6.2 \u05e4\u05e2\u05d5\u05dc\u05d5\u05ea \u05db\u05ea\u05d9\u05d1\u05d4 (\u05ea\u05dc\u05d9\u05d0\u05d6 \u2192 \u05e7\u05dc\u05d9\u05e7\u05e1)', level=2, color=BRAND_BLUE)
    add_table_rtl(doc, ['\u05e2\u05d3\u05d9\u05e4\u05d5\u05ea', '\u05e0\u05ea\u05d5\u05df', '\u05de\u05e1\u05da \u05d1\u05e7\u05dc\u05d9\u05e7\u05e1', '\u05e9\u05d9\u05d8\u05d4 \u05e0\u05d5\u05db\u05d7\u05d9\u05ea'], [
        ['**\u05e7\u05e8\u05d9\u05d8\u05d9\u05ea**', '\u05de\u05db\u05ea\u05d1 \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5', '\u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5', 'copy-paste \u05de-Google Docs'],
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05e8\u05d9\u05e9\u05d5\u05dd \u05d1\u05d9\u05e7\u05d5\u05e8', '\u05e8\u05d9\u05e9\u05d5\u05dd \u05d1\u05d9\u05e7\u05d5\u05e8\u05d9\u05dd', '\u05d4\u05d6\u05e0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05d0\u05d1\u05d7\u05e0\u05d4', '\u05d0\u05d1\u05d7\u05e0\u05d4', '\u05d7\u05d9\u05e4\u05d5\u05e9 ICD \u05d9\u05d3\u05e0\u05d9'],
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05ea\u05d5\u05db\u05e0\u05d9\u05ea \u05d8\u05d9\u05e4\u05d5\u05dc', '\u05ea\u05d5\u05db\u05e0\u05d9\u05ea \u05d8\u05d9\u05e4\u05d5\u05dc', '\u05d4\u05d6\u05e0\u05ea \u05d8\u05e7\u05e1\u05d8 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05de\u05e8\u05e9\u05de\u05d9\u05dd \u05d7\u05d3\u05e9\u05d9\u05dd', '\u05ea\u05e8\u05d5\u05e4\u05d5\u05ea \u05e7\u05d1\u05d5\u05e2\u05d5\u05ea', '\u05d7\u05d9\u05e4\u05d5\u05e9 \u05ea\u05e8\u05d5\u05e4\u05d4 \u05d9\u05d3\u05e0\u05d9'],
        ['**\u05d2\u05d1\u05d5\u05d4\u05d4**', '\u05e2\u05d3\u05db\u05d5\u05e0\u05d9 \u05de\u05d9\u05e0\u05d5\u05df', '\u05e2\u05d3\u05db\u05d5\u05df \u05de\u05d9\u05e0\u05d5\u05df', '\u05d4\u05d6\u05e0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['\u05d1\u05d9\u05e0\u05d5\u05e0\u05d9\u05ea', '\u05ea\u05d6\u05de\u05d5\u05df \u05de\u05e2\u05e7\u05d1', '\u05de\u05d5\u05d6\u05de\u05df \u05dc\u05de\u05e2\u05e7\u05d1', '\u05d1\u05d7\u05d9\u05e8\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea'],
        ['\u05d1\u05d9\u05e0\u05d5\u05e0\u05d9\u05ea', '\u05d7\u05ea\u05d9\u05de\u05d4 \u05d3\u05d9\u05d2\u05d9\u05d8\u05dc\u05d9\u05ea', '\u05d7\u05ea\u05d9\u05de\u05d4 \u05d3\u05d9\u05d2\u05d9\u05d8\u05dc\u05d9\u05ea', '\u05d4\u05d6\u05d3\u05d4\u05d5\u05ea \u05d9\u05d3\u05e0\u05d9\u05ea'],
    ])


    # ==================== SECTION 7: SECURITY ====================
    doc.add_page_break()
    add_heading_rtl(doc, '7. אבטחת מידע וגישה מרחוק', level=1)
    add_paragraph_rtl(doc, 'נושא האבטחה עלה כסוגיה מרכזית בישיבה, במיוחד בהקשר של גישה מרחוק ממיקומים מחוץ לישראל וההשלכות על אינטגרציית API.')
    add_heading_rtl(doc, '7.1 גישה מרחוק — המצב הנוכחי', level=2, color=BRAND_BLUE)
    add_quote_box(doc, 'שלומי (מכבי) — מהישיבה:', '"לחו״ל, כדי לעבוד במכבי צריך אישור מיוחד [...] זה בדרך כלל יהיה מחשב של מכבי, מחשב נייד של מכבי שהוא מנועל [...] אנחנו לא מאפשרים את זה, בטח לא בזמן מלחמה."')
    add_paragraph_rtl(doc, 'הפסיכיאטרים של תליאז עובדים מרחוק (טלפסיכיאטריה) ולעיתים שוהים בחו״ל. הגישה הנוכחית לקליקס היא דרך Citrix ודורשת:')
    add_bullet_list(doc, [
        '**Citrix StoreFront:** גישה לקליקס דרך שולחן עבודה מרוחק של מכבי',
        '**מחשב נעול:** מחשב נייד של מכבי עם הגדרות אבטחה — "מחשב שהוא מנועל, שהוא לא יכול לעשות מה שהוא רוצה במחשב"',
        '**אישור אבטחת מידע:** נדרש אישור מיוחד לגישה מחוץ לישראל',
        '**הגבלת זמן מלחמה:** שלומי ציין שבזמן מלחמה ההגבלות מחמירות אף יותר',
    ])
    add_quote_box(doc, 'דקל תליאז (מנכ"ל תליאז) — מהישיבה:', '"כי יש לנו רופאים שעובדים מחו״ל ואז בהקשר הזה אם אנחנו רוצים — הם יוכלו להוציא מרשמים וכל זה — הם לא יכולים."')
    add_heading_rtl(doc, 'פתרון מוצע — Remote Desktop דרך מחשב בארץ', level=3)
    add_quote_box(doc, 'שחר שגב (תליאז) — מהישיבה:', '"VPN או מחשב פה שהוא ייכנס למחשב בארץ ומשם ייכנס לקליקס."')
    add_quote_box(doc, 'שלומי (מכבי) — מהישיבה:', '"זה יכול להיות שמחשב בארץ אפשר להלווין אותו אם זה מחשב מנועל [...] לדעתי אני לא רואה סיבה שלא, אבל צריך לבדוק את זה. שחר, תכתוב לי בסיכום של הפגישה ואני אבדוק את זה מול אבטחת מידע."')
    add_warning_box(doc, '\u26a0 אישור CISO נדרש:', ['שלומי הבהיר שלמרות שכל הטכנולוגיה יושבת תחתיו, צוות אבטחת המידע (CISO) הוא גוף עצמאי שמקבל החלטות באופן עצמאי: "אומנם זה תחתיי, אבל הם גוף עצמאי, הם חולטים." שחר ביקש לקבל אישור פורמלי לפני יישום הפתרון.'])
    add_heading_rtl(doc, '7.2 השלכות על אינטגרציית API', level=2, color=BRAND_BLUE)
    add_info_box(doc, 'יתרון מפתח:', ['אינטגרציה באמצעות API server-to-server מבטלת לחלוטין את בעיית הגישה מחו״ל. הפסיכיאטר ממשיך לעבוד בפלטפורמת תליאז מכל מקום בעולם, והנתונים מסונכרנים ברקע ישירות בין השרתים — ללא צורך ב-VPN, Citrix, או מחשב נעול.'])
    add_paragraph_rtl(doc, 'כלומר, אינטגרציית API פותרת גם את בעיית הגישה מחו״ל וגם את בעיית ההזנה הידנית — שתי ציפורים במכה אחת.')
    add_heading_rtl(doc, '7.3 סביבת בדיקות', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, 'בישיבה, נבדקה האפשרות לקבל גישה לסביבת בדיקות של קליקס:')
    add_warning_box(doc, '\u26a0 אילוץ:', ['שלומי הבהיר שסביבת הבדיקות (TEST) של קליקס מכילה מידע אמיתי של מטופלים ולכן הגישה מוגבלת לעובדי מכבי עם הרשאות ספציפיות בלבד. לא ניתן לתת גישה חיצונית לסביבת ה-TEST. לכן, ד"ר ראמי ח׳ורי יכין את צילומי המסך מתוך המערכת הפעילה (ללא שמירת שינויים).'])
    add_paragraph_rtl(doc, 'דרישות עתידיות לסביבת בדיקות ייעודית (לאחר אישור ארכיטקט):')
    add_bullet_list(doc, [
        'סביבת sandbox ייעודית עם נתוני בדיקה פיקטיביים',
        'חשבון שירות (Service Account) לתליאז',
        'תיעוד API (endpoints, authentication, payload formats)',
    ])


    # ==================== SECTION 8: iFeel ====================
    doc.add_page_break()
    add_heading_rtl(doc, '8. אינטגרציית iFeel וסיכומים פסיכיאטריים קודמים', level=1)
    add_paragraph_rtl(doc, 'שני נושאים קריטיים שעלו בישיבה נוגעים למעבר מידע קליני ממכבי לתליאז לפני תחילת הטיפול — קבלת סיכום iFeel וקבלת סיכומים פסיכיאטריים קודמים.')
    add_heading_rtl(doc, '8.1 מהי מערכת iFeel?', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, 'iFeel היא מערכת אינטייק (מיון ראשוני) של מכבי בתחום בריאות הנפש. המטופל עובר תהליך הערכה ראשונית דרך iFeel, וסיכום ההערכה נשמר בתיק המטופל בקליקס.')
    add_quote_box(doc, 'דקל תליאז (מנכ"ל תליאז) — מהישיבה:', '"דיברנו על איך לחסוך את הזמן למטופלים, שלא יצטרכו לעבור שוב את אותם שאלונים ואותם זה. האם זה הזמן להעלות פה את הנקודה של איך אנחנו עושים פה את המעברי מידע? בין iFeel לבין הקייס מנג׳ר שלנו כדי שמטופל לא יצטרך לענות שוב על אותן שאלות."')
    add_quote_box(doc, 'יהודה (מכבי) — מהישיבה:', '"הוא בעצם מדבר על מסע המטופל החדש שבעצם יש סיכום של iFeel שיוצא ואז הוא צריך לקבל את הסיכום של iFeel כדי שלא יצטרך לעשות את כל התהליך מההתחלה אצלו."')
    add_heading_rtl(doc, 'מה תליאז צריכה מ-iFeel?', level=3)
    add_quote_box(doc, 'שחר שגב (תליאז) — מהישיבה:', '"שלומי, מבחינתי, אם יש דרך לחשוף לנו את המידע של iFeel דרך כספת או דרך API או דרך איזשהו דרך — למטופלים לפחות שאנחנו מטפלים בהם."')
    add_quote_box(doc, 'ד"ר ראמי ח׳ורי (תליאז) — מהישיבה:', '"רוצים את זה ממש בהתחלה, כשמטופל ממלא את השאלון, עוד לפני שנקבע פגישה עם מנהל תיק ופסיכיאטר. שהמערכת שלנו בעצם תתאים את המהלך של המטופל ותזהה שהוא עבר אינטייק, תזהה את המידע הקיים, ואז לחסוך ממנו לעבור אינטייק מלא."')
    add_bullet_list(doc, [
        '**מתי נדרש:** בשלב מוקדם מאוד — לפני הפגישה הראשונה, עוד לפני שנקבע פסיכיאטר',
        '**למה:** כדי למנוע מהמטופל למלא שוב שאלונים שכבר מילא ב-iFeel',
        '**איפה המידע קיים:** סיכום iFeel כבר נמצא בתיק המטופל בקליקס',
        '**מה נדרש:** דרך לחשוף את המידע הזה לתליאז — דרך API, כספת, או מנגנון אחר',
    ])
    add_info_box(doc, 'הערה:', ['שלומי ציין שנבנה כרגע "מסע מטופל חדש" סביב iFeel, אך היסס לחשוף פרטים ("אני לא יודע אם אני יכול לחשוף את זה פה"). ייתכן שהשינויים הצפויים במערכת iFeel ישפיעו על אפשרויות האינטגרציה.'])
    add_heading_rtl(doc, '8.2 סיכומים פסיכיאטריים קודמים', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, 'בנוסף ל-iFeel, תליאז צריכה גישה לסיכומים פסיכיאטריים קודמים של המטופל — מפסיכיאטרים אחרים שטיפלו בו בעבר דרך מכבי.')
    add_quote_box(doc, 'ד"ר ראמי ח׳ורי (תליאז) — מהישיבה:', '"אני אשאל את השאלה יותר רחב — גם זה יהיה מאוד מבורך אם נוכל לשלוף סיכומים פסיכיאטריים קודמים. עכשיו אם זה לא אוטומטי, אז אנחנו כן נצטרך איזושהי הרשאה של מזכירה רפואית."')
    add_quote_box(doc, 'יהודה (מכבי) — מהישיבה:', '"מזכירה רפואית לא יכולה לראות סיכומים. [...] כאילו בסוף גם מבחינה קלינית לא יכול להיות שהפסיכיאטר יראה מטופל שיש לו עבר פסיכיאטרי מבלי לראות את העבר."')
    add_critical_box(doc, 'אילוץ הרשאות:', ['מזכירות רפואיות אינן יכולות לצפות בסיכומים פסיכיאטריים (הגבלת הרשאות). אם השליפה לא תהיה אוטומטית דרך API, תידרש הרשאה של גורם רפואי מוסמך. יהודה אישר שקלינית — הפסיכיאטר חייב לראות את ההיסטוריה הפסיכיאטרית לפני תחילת הטיפול.'])
    add_table_rtl(doc, ['נתון', 'תיאור', 'עדיפות'], [
        ['סיכום iFeel', 'שאלוני אינטייק, הערכה ראשונית, מסע מטופל', '**גבוהה**'],
        ['סיכומים פסיכיאטריים קודמים', 'סיכומי ייעוץ מפסיכיאטרים קודמים', '**גבוהה**'],
        ['היסטוריית אבחנות', 'אבחנות פסיכיאטריות קודמות', 'גבוהה'],
        ['היסטוריית תרופות', 'תרופות פסיכיאטריות קודמות', 'גבוהה'],
    ])


    # ==================== SECTION 9: DATA PORTABILITY & FHIR ====================
    doc.add_page_break()
    add_heading_rtl(doc, '9. \u05d7\u05d5\u05e7 \u05e0\u05d9\u05d5\u05d3 \u05d4\u05de\u05d9\u05d3\u05e2 \u05d5-FHIR', level=1)
    add_paragraph_rtl(doc, '\u05d1\u05d9\u05e9\u05d9\u05d1\u05d4 \u05e0\u05d3\u05d5\u05e0\u05d5 \u05d4\u05d4\u05e9\u05dc\u05db\u05d5\u05ea \u05e9\u05dc \u05d7\u05d5\u05e7 \u05e0\u05d9\u05d5\u05d3 \u05d4\u05de\u05d9\u05d3\u05e2 \u05d4\u05e8\u05e4\u05d5\u05d0\u05d9 \u05d5\u05d4\u05de\u05e2\u05d1\u05e8 \u05dc\u05ea\u05e7\u05df FHIR \u05e2\u05dc \u05d0\u05e4\u05e9\u05e8\u05d5\u05d9\u05d5\u05ea \u05d4\u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4 \u05d4\u05e2\u05ea\u05d9\u05d3\u05d9\u05d5\u05ea.')
    add_quote_box(doc, '\u05e9\u05dc\u05d5\u05de\u05d9 (\u05de\u05db\u05d1\u05d9):', '"\u05e2\u05d3 2027 \u05db\u05dc \u05d4\u05e7\u05d5\u05e4\u05d5\u05ea \u05e6\u05e8\u05d9\u05db\u05d5\u05ea \u05dc\u05e2\u05d1\u05d5\u05e8 \u05dc-FHIR. \u05d6\u05d4 \u05d0\u05d5\u05de\u05e8 \u05e9\u05d1\u05e2\u05ea\u05d9\u05d3, \u05d4\u05d2\u05d9\u05e9\u05d4 \u05dc\u05de\u05d9\u05d3\u05e2 \u05d4\u05e8\u05e4\u05d5\u05d0\u05d9 \u05ea\u05d4\u05d9\u05d4 \u05e1\u05d8\u05e0\u05d3\u05e8\u05d8\u05d9\u05ea \u05d4\u05e8\u05d1\u05d4 \u05d9\u05d5\u05ea\u05e8. \u05d0\u05d1\u05dc \u05e2\u05d3 \u05d0\u05d6, \u05d0\u05e0\u05d7\u05e0\u05d5 \u05e2\u05d5\u05d1\u05d3\u05d9\u05dd \u05e2\u05dd \u05de\u05d4 \u05e9\u05d9\u05e9."')
    add_heading_rtl(doc, '9.1 \u05d7\u05d5\u05e7 \u05e0\u05d9\u05d5\u05d3 \u05d4\u05de\u05d9\u05d3\u05e2 \u05d4\u05e8\u05e4\u05d5\u05d0\u05d9', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, '\u05d7\u05d5\u05e7 \u05e0\u05d9\u05d5\u05d3 \u05d4\u05de\u05d9\u05d3\u05e2 \u05d4\u05e8\u05e4\u05d5\u05d0\u05d9 \u05de\u05d7\u05d9\u05d9\u05d1 \u05d0\u05ea \u05e7\u05d5\u05e4\u05d5\u05ea \u05d4\u05d7\u05d5\u05dc\u05d9\u05dd \u05dc\u05d0\u05e4\u05e9\u05e8 \u05d4\u05e2\u05d1\u05e8\u05ea \u05de\u05d9\u05d3\u05e2 \u05e8\u05e4\u05d5\u05d0\u05d9 \u05d1\u05e4\u05d5\u05e8\u05de\u05d8 \u05de\u05ea\u05d5\u05e7\u05e0\u05df. \u05d4\u05d4\u05e9\u05dc\u05db\u05d5\u05ea \u05e2\u05dc \u05ea\u05dc\u05d9\u05d0\u05d6:')
    add_bullet_list(doc, [
        '\u05d1\u05e2\u05ea\u05d9\u05d3, \u05d2\u05d9\u05e9\u05d4 \u05dc\u05e0\u05ea\u05d5\u05e0\u05d9 \u05de\u05d8\u05d5\u05e4\u05dc \u05ea\u05d4\u05d9\u05d4 \u05d3\u05e8\u05da API \u05e1\u05d8\u05e0\u05d3\u05e8\u05d8\u05d9 (FHIR)',
        '\u05dc\u05d0 \u05e0\u05d3\u05e8\u05e9\u05ea \u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4 \u05d9\u05d9\u05e2\u05d5\u05d3\u05d9\u05ea \u05dc\u05db\u05dc \u05e7\u05d5\u05e4\u05d4 \u2014 \u05ea\u05e7\u05df \u05d0\u05d7\u05d9\u05d3',
        '\u05d0\u05da \u05dc\u05d5\u05d7 \u05d4\u05d6\u05de\u05e0\u05d9\u05dd (2027) \u05e8\u05d7\u05d5\u05e7 \u05de\u05db\u05d3\u05d9 \u05dc\u05d4\u05de\u05ea\u05d9\u05df',
    ])
    add_heading_rtl(doc, '9.2 FHIR \u2014 Fast Healthcare Interoperability Resources', level=2, color=BRAND_BLUE)
    add_table_rtl(doc, ['\u05d4\u05d9\u05d1\u05d8', '\u05de\u05e6\u05d1 \u05e0\u05d5\u05db\u05d7\u05d9', '\u05de\u05e6\u05d1 \u05d9\u05e2\u05d3 (FHIR)'], [
        ['\u05e4\u05d5\u05e8\u05de\u05d8 \u05e0\u05ea\u05d5\u05e0\u05d9\u05dd', '\u05e4\u05e8\u05d5\u05e4\u05e8\u05d9\u05d0\u05d8\u05e8\u05d9 (\u05e7\u05dc\u05d9\u05e7\u05e1)', 'JSON/XML \u05e1\u05d8\u05e0\u05d3\u05e8\u05d8\u05d9'],
        ['Authentication', 'Citrix + \u05e4\u05e0\u05d9\u05de\u05d9', 'OAuth 2.0 / SMART on FHIR'],
        ['Resources', '\u05d8\u05d1\u05dc\u05d0\u05d5\u05ea \u05e4\u05e0\u05d9\u05de\u05d9\u05d5\u05ea', 'Patient, Encounter, Medication, Observation'],
        ['\u05ea\u05d9\u05e2\u05d5\u05d3', '\u05e4\u05e0\u05d9\u05de\u05d9 \u05d1\u05dc\u05d1\u05d3', '\u05ea\u05e7\u05df \u05e4\u05d5\u05de\u05d1\u05d9'],
        ['\u05d6\u05de\u05d9\u05e0\u05d5\u05ea', '\u05d4\u05d9\u05d5\u05dd (\u05e0\u05d3\u05e8\u05e9 API \u05d9\u05d9\u05e2\u05d5\u05d3\u05d9)', '2027 (\u05dc\u05e4\u05d9 \u05d4\u05e8\u05d2\u05d5\u05dc\u05e6\u05d9\u05d4)'],
    ])
    add_heading_rtl(doc, '9.3 \u05d0\u05e1\u05d8\u05e8\u05d8\u05d2\u05d9\u05d4 \u05de\u05d5\u05de\u05dc\u05e6\u05ea', level=2, color=BRAND_BLUE)
    add_info_box(doc, '\u05d2\u05d9\u05e9\u05d4 \u05d3\u05d5-\u05e9\u05dc\u05d1\u05d9\u05ea:', [
        '1. **\u05dc\u05d8\u05d5\u05d5\u05d7 \u05e7\u05e6\u05e8 (2026):** \u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4 \u05d1\u05d0\u05de\u05e6\u05e2\u05d5\u05ea API \u05d9\u05d9\u05e2\u05d5\u05d3\u05d9 \u05de\u05d5\u05dc \u05e7\u05dc\u05d9\u05e7\u05e1 (\u05db\u05de\u05ea\u05d5\u05d0\u05e8 \u05d1\u05de\u05e1\u05de\u05da \u05d6\u05d4)',
        '2. **\u05dc\u05d8\u05d5\u05d5\u05d7 \u05d0\u05e8\u05d5\u05da (2027+):** \u05de\u05e2\u05d1\u05e8 \u05dc-FHIR \u05db\u05e9\u05d4\u05ea\u05e7\u05df \u05d9\u05d4\u05d9\u05d4 \u05d6\u05de\u05d9\u05df, \u05ea\u05d5\u05da \u05e9\u05de\u05d9\u05e8\u05d4 \u05e2\u05dc \u05ea\u05d0\u05d9\u05de\u05d5\u05ea \u05dc\u05d0\u05d7\u05d5\u05e8',
    ])


    # ==================== SECTION 10: MENTAL HEALTH PARTITION ====================
    doc.add_page_break()
    add_heading_rtl(doc, '10. \u05de\u05d7\u05d9\u05e6\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9 \u2014 \u05d7\u05e1\u05d9\u05de\u05ea \u05de\u05d9\u05d3\u05e2 \u05de\u05e1\u05e4\u05e7\u05d9\u05dd \u05d7\u05d9\u05e6\u05d5\u05e0\u05d9\u05d9\u05dd', level=1)
    add_paragraph_rtl(doc, '\u05d0\u05d7\u05d3 \u05d4\u05e0\u05d5\u05e9\u05d0\u05d9\u05dd \u05d4\u05de\u05e9\u05de\u05e2\u05d5\u05ea\u05d9\u05d9\u05dd \u05d1\u05d9\u05d5\u05ea\u05e8 \u05e9\u05e2\u05dc\u05d5 \u05dc\u05d0\u05d7\u05e8 \u05d9\u05e9\u05d9\u05d1\u05ea \u05d4\u05ea\u05d9\u05d0\u05d5\u05dd \u05d4\u05d5\u05d0 \u05e0\u05d5\u05e9\u05d0 \u05de\u05d7\u05d9\u05e6\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9 \u05d1\u05de\u05e2\u05e8\u05db\u05ea \u05e7\u05dc\u05d9\u05e7\u05e1. \u05e2\u05e0\u05ea \u05de\u05d9\u05e8\u05d5\u05df, \u05de\u05e0\u05d4\u05dc\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9 \u05d1\u05de\u05db\u05d1\u05d9, \u05d4\u05d1\u05d4\u05d9\u05e8\u05d4 \u05d1\u05d0\u05d9\u05de\u05d9\u05d9\u05dc \u05de-20 \u05d1\u05d0\u05e4\u05e8\u05d9\u05dc 2025 \u05e9\u05de\u05d9\u05d3\u05e2 \u05e4\u05e1\u05d9\u05db\u05d9\u05d0\u05d8\u05e8\u05d9 \u05de\u05e1\u05e4\u05e7\u05d9\u05dd \u05d7\u05d9\u05e6\u05d5\u05e0\u05d9\u05d9\u05dd \u05d7\u05e1\u05d5\u05dd \u05dc\u05d7\u05dc\u05d5\u05d8\u05d9\u05df \u05de\u05e8\u05d5\u05e4\u05d0\u05d9 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4.')
    add_quote_box(doc, '\u05e2\u05e0\u05ea \u05de\u05d9\u05e8\u05d5\u05df (\u05de\u05e0\u05d4\u05dc\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9, \u05de\u05db\u05d1\u05d9) \u2014 \u05d0\u05d9\u05de\u05d9\u05d9\u05dc, 20 \u05d1\u05d0\u05e4\u05e8\u05d9\u05dc 2025:', '"\u05d9\u05e9\u05e0\u05d4 \u05de\u05d7\u05d9\u05e6\u05d4 \u05e9\u05dc \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9. \u05e8\u05d5\u05e4\u05d0 \u05e8\u05d0\u05e9\u05d5\u05e0\u05d9 \u05dc\u05d0 \u05d9\u05db\u05d5\u05dc \u05dc\u05e8\u05d0\u05d5\u05ea \u05de\u05d9\u05d3\u05e2 \u05e9\u05dc \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9 \u05e9\u05dc \u05e1\u05e4\u05e7\u05d9\u05dd \u05d7\u05d9\u05e6\u05d5\u05e0\u05d9\u05d9\u05dd. [...] \u05db\u05e9\u05d4\u05e1\u05e4\u05e7 \u05d4\u05d7\u05d9\u05e6\u05d5\u05e0\u05d9 \u05e9\u05d5\u05dc\u05d7 \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5, \u05d4\u05d9\u05d0 \u05dc\u05d0 \u05de\u05d2\u05d9\u05e2\u05d4 \u05dc\u05e8\u05d5\u05e4\u05d0 \u05d0\u05dc\u05d0 \u05e8\u05e7 \u05dc\u05de\u05e8\u05e4\u05d0\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9 \u05d4\u05de\u05d7\u05d5\u05d6\u05d9\u05ea."')
    add_quote_box(doc, '\u05d3\u05e7\u05dc \u05ea\u05dc\u05d9\u05d0\u05d6 (\u05de\u05e0\u05db\u05f4\u05dc \u05ea\u05dc\u05d9\u05d0\u05d6) \u2014 \u05de\u05d4\u05d9\u05e9\u05d9\u05d1\u05d4:', '"\u05d0\u05dd \u05d4\u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05ea\u05d9 \u05dc\u05d0 \u05e8\u05d5\u05d0\u05d4 \u05d0\u05ea \u05d4\u05e1\u05d9\u05db\u05d5\u05dd \u05e9\u05dc\u05e0\u05d5, \u05d6\u05d4 \u05de\u05d1\u05d8\u05dc \u05d0\u05ea \u05db\u05dc \u05d4\u05e2\u05e8\u05da \u05e9\u05dc \u05d4\u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4. \u05d7\u05d9\u05d9\u05d1\u05d9\u05dd \u05dc\u05de\u05e6\u05d5\u05d0 \u05e4\u05ea\u05e8\u05d5\u05df \u05dc\u05d6\u05d4."')
    add_quote_box(doc, '\u05e9\u05d7\u05e8 \u05e9\u05d2\u05d1 (\u05de\u05e0\u05d4\u05dc\u05ea \u05de\u05e8\u05e4\u05d0\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9, \u05de\u05db\u05d1\u05d9 \u05de\u05e8\u05db\u05d6) \u2014 \u05de\u05d4\u05d9\u05e9\u05d9\u05d1\u05d4:', '"\u05d4\u05de\u05d7\u05d9\u05e6\u05d4 \u05d4\u05d9\u05d0 \u05de\u05db\u05d5\u05d5\u05e0\u05ea \u2014 \u05d4\u05d9\u05d0 \u05de\u05d2\u05d9\u05e0\u05d4 \u05e2\u05dc \u05e4\u05e8\u05d8\u05d9\u05d5\u05ea \u05d4\u05de\u05d8\u05d5\u05e4\u05dc. \u05d0\u05d1\u05dc \u05d9\u05e9 \u05de\u05e7\u05e8\u05d9\u05dd \u05e9\u05d1\u05d4\u05dd \u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4 \u05e6\u05e8\u05d9\u05da \u05dc\u05d3\u05e2\u05ea \u05e2\u05dc \u05d4\u05d8\u05d9\u05e4\u05d5\u05dc \u05d4\u05e4\u05e1\u05d9\u05db\u05d9\u05d0\u05d8\u05e8\u05d9, \u05d5\u05d0\u05d6 \u05d0\u05e0\u05d7\u05e0\u05d5 \u05de\u05e2\u05d1\u05d9\u05e8\u05d9\u05dd \u05d0\u05ea \u05d4\u05de\u05d9\u05d3\u05e2 \u05d9\u05d3\u05e0\u05d9\u05ea."')
    add_critical_box(doc, '\u05d4\u05e9\u05dc\u05db\u05d4 \u05e7\u05e8\u05d9\u05d8\u05d9\u05ea \u05e2\u05dc \u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d9\u05ea API:', ['\u05d2\u05dd \u05d0\u05dd \u05ea\u05dc\u05d9\u05d0\u05d6 \u05ea\u05e2\u05dc\u05d4 \u05e1\u05d9\u05db\u05d5\u05de\u05d9\u05dd \u05d1\u05d0\u05de\u05e6\u05e2\u05d5\u05ea API \u05d9\u05e9\u05d9\u05e8\u05d5\u05ea \u05dc\u05e7\u05dc\u05d9\u05e7\u05e1, \u05d9\u05d9\u05ea\u05db\u05df \u05e9\u05d4\u05de\u05d7\u05d9\u05e6\u05d4 \u05ea\u05d7\u05e1\u05d5\u05dd \u05d0\u05ea \u05d4\u05d2\u05d9\u05e9\u05d4 \u05e9\u05dc \u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4 \u05dc\u05de\u05d9\u05d3\u05e2 \u05d6\u05d4. \u05d9\u05e9 \u05dc\u05d1\u05e8\u05e8: (\u05d0) \u05d4\u05d0\u05dd \u05e0\u05d9\u05ea\u05df \u05dc\u05e9\u05dc\u05d5\u05d7 \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5 \u05e9\u05ea\u05e2\u05e7\u05d5\u05e3 \u05d0\u05ea \u05d4\u05de\u05d7\u05d9\u05e6\u05d4, (\u05d1) \u05d4\u05d0\u05dd \u05e0\u05d3\u05e8\u05e9 \u05d0\u05d9\u05e9\u05d5\u05e8 \u05e9\u05dc \u05de\u05e8\u05e4\u05d0\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9 \u05d4\u05de\u05d7\u05d5\u05d6\u05d9\u05ea, (\u05d2) \u05d4\u05d0\u05dd \u05d4\u05de\u05d8\u05d5\u05e4\u05dc \u05d9\u05db\u05d5\u05dc \u05dc\u05ea\u05ea \u05d4\u05e1\u05db\u05de\u05d4 \u05dc\u05e9\u05d9\u05ea\u05d5\u05e3 \u05d4\u05de\u05d9\u05d3\u05e2.'])
    add_heading_rtl(doc, '10.1 \u05d4\u05e9\u05dc\u05db\u05d5\u05ea \u05e2\u05dc \u05d0\u05e1\u05d8\u05e8\u05d8\u05d2\u05d9\u05d9\u05ea \u05d4\u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, '\u05e0\u05d5\u05e9\u05d0 \u05d4\u05de\u05d7\u05d9\u05e6\u05d4 \u05de\u05e9\u05e0\u05d4 \u05d0\u05ea \u05e1\u05d3\u05e8 \u05d4\u05e2\u05d3\u05d9\u05e4\u05d5\u05d9\u05d5\u05ea \u05e9\u05dc \u05d4\u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4:')
    add_bullet_list(doc, [
        '**\u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5:** \u05d9\u05e9 \u05dc\u05d5\u05d5\u05d3\u05d0 \u05e9\u05d4-API \u05d9\u05d5\u05d3\u05e2 \u05dc\u05e0\u05ea\u05d1 \u05d0\u05ea \u05d4\u05ea\u05e9\u05d5\u05d1\u05d4 \u05d2\u05dd \u05dc\u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4 \u05d5\u05d2\u05dd \u05dc\u05de\u05e8\u05e4\u05d0\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9',
        '**\u05d4\u05e8\u05e9\u05d0\u05d5\u05ea:** \u05d9\u05d9\u05ea\u05db\u05df \u05e9\u05e0\u05d3\u05e8\u05e9\u05ea \u05d4\u05e8\u05e9\u05d0\u05d4 \u05de\u05d9\u05d5\u05d7\u05d3\u05ea \u05db\u05d3\u05d9 \u05dc\u05e2\u05e7\u05d5\u05e3 \u05d0\u05ea \u05d4\u05de\u05d7\u05d9\u05e6\u05d4',
        '**\u05d4\u05e1\u05db\u05de\u05ea \u05de\u05d8\u05d5\u05e4\u05dc:** \u05d9\u05d9\u05ea\u05db\u05df \u05e9\u05e0\u05d3\u05e8\u05e9 \u05de\u05e0\u05d2\u05e0\u05d5\u05df \u05d4\u05e1\u05db\u05de\u05d4 \u05d3\u05d9\u05d2\u05d9\u05d8\u05dc\u05d9 \u05e9\u05dc \u05d4\u05de\u05d8\u05d5\u05e4\u05dc \u05dc\u05e9\u05d9\u05ea\u05d5\u05e3 \u05de\u05d9\u05d3\u05e2 \u05e4\u05e1\u05d9\u05db\u05d9\u05d0\u05d8\u05e8\u05d9 \u05e2\u05dd \u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4',
    ])
    add_decision_box(doc, '\u05de\u05e1\u05e7\u05e0\u05d4 \u05d0\u05e1\u05d8\u05e8\u05d8\u05d2\u05d9\u05ea:', ['\u05e0\u05d5\u05e9\u05d0 \u05d4\u05de\u05d7\u05d9\u05e6\u05d4 \u05d4\u05d5\u05d0 \u05d7\u05e1\u05dd \u05e4\u05d5\u05d8\u05e0\u05e6\u05d9\u05d0\u05dc\u05d9 \u05de\u05e8\u05db\u05d6\u05d9. \u05d9\u05e9 \u05dc\u05e7\u05d1\u05d5\u05e2 \u05e4\u05d2\u05d9\u05e9\u05d4 \u05d9\u05d9\u05e2\u05d5\u05d3\u05d9\u05ea \u05e2\u05dd \u05e2\u05e0\u05ea \u05de\u05d9\u05e8\u05d5\u05df \u05d5\u05e6\u05d5\u05d5\u05ea \u05d4\u05d0\u05d3\u05e8\u05d9\u05db\u05dc\u05d5\u05ea \u05e9\u05dc \u05de\u05db\u05d1\u05d9 \u05dc\u05e4\u05e0\u05d9 \u05ea\u05d7\u05d9\u05dc\u05ea \u05d4\u05e4\u05d9\u05ea\u05d5\u05d7, \u05db\u05d3\u05d9 \u05dc\u05de\u05e4\u05d5\u05ea \u05d0\u05ea \u05db\u05dc\u05dc\u05d9 \u05d4\u05de\u05d7\u05d9\u05e6\u05d4 \u05d5\u05dc\u05ea\u05db\u05e0\u05df \u05d0\u05ea \u05d4-API \u05d1\u05d4\u05ea\u05d0\u05dd.'])


    # ==================== SECTION 11: TEST CASES & EMAIL CHAIN ====================
    doc.add_page_break()
    add_heading_rtl(doc, '11. \u05de\u05e7\u05e8\u05d9 \u05d1\u05d3\u05d9\u05e7\u05d4 \u05d5\u05e9\u05e8\u05e9\u05e8\u05ea \u05d0\u05d9\u05de\u05d9\u05d9\u05dc\u05d9\u05dd \u2014 \u05d4\u05ea\u05e0\u05e1\u05d5\u05ea \u05e2\u05dd \u05d4\u05db\u05e1\u05e4\u05ea', level=1)
    add_paragraph_rtl(doc, '\u05d1\u05de\u05d4\u05dc\u05da \u05de\u05e8\u05e5-\u05d0\u05e4\u05e8\u05d9\u05dc 2025, \u05ea\u05dc\u05d9\u05d0\u05d6 \u05d1\u05d9\u05e6\u05e2\u05d4 \u05e1\u05d3\u05e8\u05ea \u05de\u05e7\u05e8\u05d9 \u05d1\u05d3\u05d9\u05e7\u05d4 \u05e2\u05dd \u05de\u05e0\u05d2\u05e0\u05d5\u05df \u05d4\u05db\u05e1\u05e4\u05ea \u05db\u05d3\u05d9 \u05dc\u05d1\u05d7\u05d5\u05df \u05d0\u05ea \u05d4\u05d9\u05ea\u05db\u05e0\u05d5\u05ea \u05d4\u05e2\u05dc\u05d0\u05ea \u05e1\u05d9\u05db\u05d5\u05de\u05d9\u05dd \u05d0\u05d5\u05d8\u05d5\u05de\u05d8\u05d9\u05ea. \u05dc\u05d4\u05dc\u05df \u05ea\u05d5\u05e6\u05d0\u05d5\u05ea \u05d4\u05d1\u05d3\u05d9\u05e7\u05d5\u05ea:')
    add_table_rtl(doc, ['\u05de\u05e1\u05f3', '\u05ea.\u05d6. (\u05d7\u05dc\u05e7\u05d9)', '\u05ea\u05d9\u05d0\u05d5\u05e8', '\u05e9\u05d0\u05dc\u05d4 \u05e9\u05e0\u05d1\u05d3\u05e7\u05d4', '\u05ea\u05d5\u05e6\u05d0\u05d4'], [
        ['1', '***4567', '\u05e1\u05d9\u05db\u05d5\u05dd \u05d0\u05d9\u05e0\u05d8\u05d9\u05d9\u05e7 \u2014 \u05d4\u05e2\u05dc\u05d0\u05d4 \u05e8\u05d0\u05e9\u05d5\u05e0\u05d4', '\u05d4\u05d0\u05dd \u05d4\u05de\u05e1\u05de\u05da \u05de\u05d2\u05d9\u05e2 \u05dc\u05ea\u05d9\u05e7?', '\u05d4\u05d2\u05d9\u05e2 \u05dc\u05ea\u05d9\u05e7 \u05d0\u05da \u05dc\u05d0 \u05dc\u05d3\u05d5\u05d0\u05e8 \u05e8\u05e4\u05d5\u05d0\u05d9'],
        ['2', '***4567', '\u05e1\u05d9\u05db\u05d5\u05dd \u05de\u05e2\u05e7\u05d1 \u2014 \u05d0\u05d5\u05ea\u05d5 \u05de\u05d8\u05d5\u05e4\u05dc', '\u05d4\u05d0\u05dd \u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4 \u05e8\u05d5\u05d0\u05d4?', '\u05dc\u05d0 \u2014 \u05d4\u05de\u05e1\u05de\u05da \u05d1\u05ea\u05d9\u05e7 \u05d0\u05da \u05dc\u05dc\u05d0 \u05d4\u05ea\u05e8\u05d0\u05d4'],
        ['3', '***8901', '\u05de\u05d8\u05d5\u05e4\u05dc \u05d7\u05d3\u05e9 \u2014 \u05d0\u05d9\u05e0\u05d8\u05d9\u05d9\u05e7', '\u05d4\u05d0\u05dd \u05e0\u05d5\u05e6\u05e8\u05ea \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5?', '\u05dc\u05d0 \u2014 \u05d4\u05db\u05e1\u05e4\u05ea \u05de\u05e2\u05dc\u05d4 \u05e7\u05d5\u05d1\u05e5, \u05dc\u05d0 \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5'],
        ['4', '***2345', '\u05e1\u05d9\u05db\u05d5\u05dd \u05e2\u05dd \u05e7\u05d5\u05d3 \u05d0\u05d1\u05d7\u05e0\u05d4', '\u05d4\u05d0\u05dd ICD \u05e0\u05e7\u05dc\u05d8?', '\u05dc\u05d0 \u2014 \u05d4\u05db\u05e1\u05e4\u05ea \u05dc\u05d0 \u05ea\u05d5\u05de\u05db\u05ea \u05d1\u05e9\u05d3\u05d5\u05ea \u05de\u05d5\u05d1\u05e0\u05d9\u05dd'],
    ])
    add_heading_rtl(doc, '11.1 \u05e9\u05e8\u05e9\u05e8\u05ea \u05d0\u05d9\u05de\u05d9\u05d9\u05dc\u05d9\u05dd \u2014 \u05e6\u05d9\u05e8 \u05d6\u05de\u05df', level=2, color=BRAND_BLUE)
    add_paragraph_rtl(doc, '\u05dc\u05d4\u05dc\u05df \u05e6\u05d9\u05e8 \u05d4\u05d6\u05de\u05df \u05d4\u05de\u05dc\u05d0 \u05e9\u05dc \u05d4\u05d4\u05ea\u05db\u05ea\u05d1\u05d5\u05ea \u05d1\u05e0\u05d5\u05e9\u05d0 \u05d4\u05db\u05e1\u05e4\u05ea \u05d5\u05d4\u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4:')
    add_table_rtl(doc, ['\u05ea\u05d0\u05e8\u05d9\u05da', '\u05e9\u05d5\u05dc\u05d7', '\u05e0\u05d5\u05e9\u05d0', '\u05ea\u05d5\u05db\u05df \u05e2\u05d9\u05e7\u05e8\u05d9'], [
        ['9/3/25', '\u05d3\u05f4\u05e8 \u05e8\u05d0\u05de\u05d9 \u05d7\u05f3\u05d5\u05e8\u05d9', '\u05e4\u05e0\u05d9\u05d9\u05d4 \u05e8\u05d0\u05e9\u05d5\u05e0\u05d9\u05ea \u05dc\u05e9\u05dc\u05d5\u05de\u05d9', '\u05d1\u05e7\u05e9\u05d4 \u05dc\u05d1\u05d7\u05d5\u05df \u05d0\u05e4\u05e9\u05e8\u05d5\u05d9\u05d5\u05ea \u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4 \u05e2\u05dd \u05e7\u05dc\u05d9\u05e7\u05e1'],
        ['12/3/25', '\u05e9\u05dc\u05d5\u05de\u05d9 (\u05de\u05db\u05d1\u05d9)', '\u05ea\u05e9\u05d5\u05d1\u05d4 \u05e8\u05d0\u05e9\u05d5\u05e0\u05d9\u05ea', '\u05d4\u05e4\u05e0\u05d9\u05d9\u05d4 \u05dc\u05db\u05e1\u05e4\u05ea \u05db\u05e4\u05ea\u05e8\u05d5\u05df \u05de\u05d9\u05d9\u05d3\u05d9 + \u05d4\u05d6\u05de\u05e0\u05d4 \u05dc\u05d9\u05e9\u05d9\u05d1\u05d4'],
        ['19/3/25', '\u05db\u05d5\u05dc\u05dd', '\u05d9\u05e9\u05d9\u05d1\u05ea \u05ea\u05d9\u05d0\u05d5\u05dd', '\u05d9\u05e9\u05d9\u05d1\u05d4 \u05e4\u05e0\u05d9\u05dd-\u05d0\u05dc-\u05e4\u05e0\u05d9\u05dd \u2014 \u05e8\u05d0\u05d5 \u05e1\u05e2\u05d9\u05e3 2'],
        ['24/3/25', '\u05d3\u05f4\u05e8 \u05e8\u05d0\u05de\u05d9 \u05d7\u05f3\u05d5\u05e8\u05d9', '\u05de\u05e1\u05de\u05da \u05de\u05e4\u05e8\u05d8 \u05d8\u05db\u05e0\u05d9 v1', '\u05d8\u05d9\u05d5\u05d8\u05d4 \u05e8\u05d0\u05e9\u05d5\u05e0\u05d4 \u05e9\u05dc \u05de\u05e1\u05de\u05da \u05d6\u05d4'],
        ['2/4/25', '\u05e9\u05dc\u05d5\u05de\u05d9 (\u05de\u05db\u05d1\u05d9)', '\u05e2\u05d3\u05db\u05d5\u05df \u05db\u05e1\u05e4\u05ea', '\u05e4\u05e8\u05d8\u05d9 \u05d2\u05d9\u05e9\u05d4 \u05dc\u05db\u05e1\u05e4\u05ea \u05dc\u05d1\u05d3\u05d9\u05e7\u05d5\u05ea'],
        ['8/4/25', '\u05d3\u05f4\u05e8 \u05e8\u05d0\u05de\u05d9 \u05d7\u05f3\u05d5\u05e8\u05d9', '\u05d3\u05d9\u05d5\u05d5\u05d7 \u05d1\u05d3\u05d9\u05e7\u05d4 #1-2', '\u05e1\u05d9\u05db\u05d5\u05de\u05d9\u05dd \u05d4\u05d5\u05e2\u05dc\u05d5 \u05d0\u05da \u05dc\u05d0 \u05d4\u05d2\u05d9\u05e2\u05d5 \u05dc\u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4'],
        ['10/4/25', '\u05e9\u05dc\u05d5\u05de\u05d9 (\u05de\u05db\u05d1\u05d9)', '\u05ea\u05e9\u05d5\u05d1\u05d4', '\u05d1\u05d3\u05d9\u05e7\u05d4 \u05e4\u05e0\u05d9\u05de\u05d9\u05ea \u2014 \u05d4\u05db\u05e1\u05e4\u05ea \u05dc\u05d0 \u05d9\u05d5\u05e6\u05e8\u05ea \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5'],
        ['15/4/25', '\u05d3\u05f4\u05e8 \u05e8\u05d0\u05de\u05d9 \u05d7\u05f3\u05d5\u05e8\u05d9', '\u05d3\u05d9\u05d5\u05d5\u05d7 \u05d1\u05d3\u05d9\u05e7\u05d4 #3-4', '\u05d0\u05d9\u05e9\u05d5\u05e8 \u05e9\u05d4\u05db\u05e1\u05e4\u05ea \u05dc\u05d0 \u05de\u05ea\u05d0\u05d9\u05de\u05d4 \u05dc\u05e6\u05e8\u05db\u05d9 \u05ea\u05dc\u05d9\u05d0\u05d6'],
        ['20/4/25', '\u05e2\u05e0\u05ea \u05de\u05d9\u05e8\u05d5\u05df (\u05de\u05db\u05d1\u05d9)', '\u05e2\u05d3\u05db\u05d5\u05df \u05de\u05d7\u05d9\u05e6\u05d4', '\u05d4\u05d1\u05d4\u05e8\u05d4 \u05e2\u05dc \u05de\u05d7\u05d9\u05e6\u05ea \u05d1\u05e8\u05d9\u05d0\u05d5\u05ea \u05d4\u05e0\u05e4\u05e9 \u2014 \u05e8\u05d0\u05d5 \u05e1\u05e2\u05d9\u05e3 10'],
        ['22/4/25', '\u05d3\u05f4\u05e8 \u05e8\u05d0\u05de\u05d9 \u05d7\u05f3\u05d5\u05e8\u05d9', '\u05de\u05e1\u05de\u05da \u05de\u05e4\u05e8\u05d8 v2', '\u05d2\u05e8\u05e1\u05d4 \u05de\u05e2\u05d5\u05d3\u05db\u05e0\u05ea \u05db\u05d5\u05dc\u05dc \u05de\u05de\u05e6\u05d0\u05d9 \u05db\u05e1\u05e4\u05ea \u05d5\u05de\u05d7\u05d9\u05e6\u05d4 (\u05de\u05e1\u05de\u05da \u05d6\u05d4)'],
    ])
    add_critical_box(doc, '\u05de\u05e1\u05e7\u05e0\u05d4:', ['\u05d4\u05db\u05e1\u05e4\u05ea \u05d0\u05d9\u05e0\u05d4 \u05de\u05ea\u05d0\u05d9\u05de\u05d4 \u05db\u05e4\u05ea\u05e8\u05d5\u05df \u05e7\u05d1\u05d5\u05e2 \u05dc\u05ea\u05dc\u05d9\u05d0\u05d6. \u05d4\u05d9\u05d0 \u05de\u05e2\u05dc\u05d4 \u05e7\u05d1\u05e6\u05d9\u05dd \u05d1\u05dc\u05d1\u05d3, \u05dc\u05dc\u05d0 \u05d9\u05e6\u05d9\u05e8\u05ea \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5, \u05dc\u05dc\u05d0 \u05e9\u05d3\u05d5\u05ea \u05de\u05d5\u05d1\u05e0\u05d9\u05dd (ICD, \u05ea\u05e8\u05d5\u05e4\u05d5\u05ea), \u05d5\u05dc\u05dc\u05d0 \u05d4\u05ea\u05e8\u05d0\u05d4 \u05dc\u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4. \u05e0\u05d3\u05e8\u05e9 API \u05d9\u05e9\u05d9\u05e8 \u05dc\u05e7\u05dc\u05d9\u05e7\u05e1 \u05db\u05e4\u05d9 \u05e9\u05de\u05ea\u05d5\u05d0\u05e8 \u05d1\u05e1\u05e2\u05d9\u05e4\u05d9\u05dd 3-6.'])


    # ==================== SECTION 12: INTEGRATION PHASES ====================
    doc.add_page_break()
    add_heading_rtl(doc, '12. \u05e9\u05dc\u05d1\u05d9 \u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4 \u05de\u05d5\u05de\u05dc\u05e6\u05d9\u05dd \u05d5\u05e4\u05e8\u05d9\u05d8\u05d9 \u05e4\u05e2\u05d5\u05dc\u05d4', level=1)
    add_paragraph_rtl(doc, '\u05e2\u05dc \u05d1\u05e1\u05d9\u05e1 \u05d4\u05de\u05de\u05e6\u05d0\u05d9\u05dd \u05de\u05e1\u05e2\u05d9\u05e4\u05d9\u05dd 1-11, \u05dc\u05d4\u05dc\u05df \u05ea\u05d5\u05db\u05e0\u05d9\u05ea \u05d4\u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d4 \u05d4\u05de\u05d5\u05de\u05dc\u05e6\u05ea \u05d1\u05d2\u05d9\u05e9\u05d4 \u05de\u05d3\u05d5\u05e8\u05d2\u05ea:')
    add_quote_box(doc, '\u05d9\u05d4\u05d5\u05d3\u05d4 (\u05e8\u05d0\u05e9 \u05d0\u05d2\u05e3 \u05de\u05e2\u05e8\u05db\u05d5\u05ea \u05de\u05d9\u05d3\u05e2, \u05de\u05db\u05d1\u05d9) \u2014 \u05de\u05d4\u05d9\u05e9\u05d9\u05d1\u05d4:', '"\u05d1\u05d5\u05d0\u05d5 \u05e0\u05e2\u05e9\u05d4 \u05d0\u05ea \u05d6\u05d4 \u05d1\u05e9\u05dc\u05d1\u05d9\u05dd. \u05e7\u05d5\u05d3\u05dd \u05db\u05dc \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5, \u05d0\u05d7\u05e8 \u05db\u05da \u05de\u05e8\u05e9\u05de\u05d9\u05dd, \u05d5\u05d1\u05e1\u05d5\u05e3 \u05e1\u05e0\u05db\u05e8\u05d5\u05df \u05de\u05dc\u05d0. \u05db\u05db\u05d4 \u05d0\u05e0\u05d7\u05e0\u05d5 \u05d9\u05db\u05d5\u05dc\u05d9\u05dd \u05dc\u05d5\u05d5\u05d3\u05d0 \u05e9\u05db\u05dc \u05e9\u05dc\u05d1 \u05e2\u05d5\u05d1\u05d3 \u05dc\u05e4\u05e0\u05d9 \u05e9\u05de\u05de\u05e9\u05d9\u05db\u05d9\u05dd."')
    add_phase_box(doc, '\u05e9\u05dc\u05d1 1 \u2014 \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5 (3-4 \u05d7\u05d5\u05d3\u05e9\u05d9\u05dd)', [
        '\u05d4\u05e2\u05dc\u05d0\u05d4 \u05d0\u05d5\u05d8\u05d5\u05de\u05d8\u05d9\u05ea \u05e9\u05dc \u05e1\u05d9\u05db\u05d5\u05dd \u05d9\u05d9\u05e2\u05d5\u05e5 \u05de\u05ea\u05dc\u05d9\u05d0\u05d6 \u05dc\u05e7\u05dc\u05d9\u05e7\u05e1 \u05db\u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5 \u05e8\u05e9\u05de\u05d9\u05ea.',
        '\u05db\u05d5\u05dc\u05dc: \u05e0\u05e8\u05d8\u05d9\u05d1 \u05e7\u05dc\u05d9\u05e0\u05d9, \u05d0\u05d1\u05d7\u05e0\u05d5\u05ea ICD, \u05d4\u05de\u05dc\u05e6\u05d5\u05ea \u05d8\u05d9\u05e4\u05d5\u05dc, \u05d5\u05e0\u05d9\u05ea\u05d5\u05d1 \u05dc\u05e8\u05d5\u05e4\u05d0 \u05d4\u05de\u05e9\u05e4\u05d7\u05d4.',
        '**\u05d6\u05d4\u05d5 \u05d4-MVP.**',
    ])
    add_phase_box_yellow(doc, '\u05e9\u05dc\u05d1 2 \u2014 \u05e7\u05e8\u05d9\u05d0\u05ea \u05e0\u05ea\u05d5\u05e0\u05d9 \u05de\u05d8\u05d5\u05e4\u05dc (2-3 \u05d7\u05d5\u05d3\u05e9\u05d9\u05dd)', [
        '\u05e9\u05dc\u05d9\u05e4\u05ea \u05d3\u05de\u05d5\u05d2\u05e8\u05e4\u05d9\u05d4, \u05d0\u05d1\u05d7\u05e0\u05d5\u05ea \u05d9\u05d3\u05d5\u05e2\u05d5\u05ea, \u05ea\u05e8\u05d5\u05e4\u05d5\u05ea \u05e0\u05d5\u05db\u05d7\u05d9\u05d5\u05ea \u05d5\u05d0\u05dc\u05e8\u05d2\u05d9\u05d5\u05ea \u05de\u05e7\u05dc\u05d9\u05e7\u05e1 \u05d0\u05dc \u05ea\u05d5\u05da \u05e4\u05dc\u05d8\u05e4\u05d5\u05e8\u05de\u05ea \u05ea\u05dc\u05d9\u05d0\u05d6.',
        '\u05de\u05e4\u05d7\u05d9\u05ea \u05d4\u05d6\u05e0\u05d4 \u05d9\u05d3\u05e0\u05d9\u05ea \u05d5\u05de\u05e9\u05e4\u05e8\u05ea \u05d1\u05d8\u05d9\u05d7\u05d5\u05ea \u05e7\u05dc\u05d9\u05e0\u05d9\u05ea.',
    ])
    add_phase_box_orange(doc, '\u05e9\u05dc\u05d1 3 \u2014 \u05de\u05e8\u05e9\u05de\u05d9\u05dd (3-4 \u05d7\u05d5\u05d3\u05e9\u05d9\u05dd)', [
        '\u05d4\u05e4\u05e7\u05ea \u05de\u05e8\u05e9\u05de\u05d9\u05dd \u05d0\u05dc\u05e7\u05d8\u05e8\u05d5\u05e0\u05d9\u05d9\u05dd \u05d1\u05d0\u05de\u05e6\u05e2\u05d5\u05ea API \u2014 \u05db\u05d5\u05dc\u05dc \u05d7\u05d9\u05e4\u05d5\u05e9 \u05ea\u05e8\u05d5\u05e4\u05d4, \u05d4\u05d2\u05d3\u05e8\u05ea \u05de\u05d9\u05e0\u05d5\u05df, \u05d5\u05d7\u05ea\u05d9\u05de\u05d4 \u05d3\u05d9\u05d2\u05d9\u05d8\u05dc\u05d9\u05ea.',
        '\u05d3\u05d5\u05e8\u05e9 \u05d0\u05d9\u05e9\u05d5\u05e8 \u05e8\u05d2\u05d5\u05dc\u05d8\u05d5\u05e8\u05d9 \u05e0\u05d5\u05e1\u05e3.',
    ])
    add_warning_box(doc, '\u26a0 מורכבות מרשמים:', ['יהודה הדגיש בישיבה שמרשם בקליקס אינו עומד בפני עצמו — הוא תוצר של ביקור (visit). כדי להוציא מרשם, יש קודם לפתוח ביקור, להזין אבחנה, ולהזין פעולה: "כדי להוציא מרשם אני חייב קודם למלא את השדות שלפני [...] מרשם מוצא כתוצר של ביקור." לכן אינטגרציית מרשמים דורשת אינטגרציה מלאה של תהליך הביקור. עם זאת, שלומי ציין: "יכול להיות שדווקא מרשמים זה משהו קל" — ולכן יש לבחון את המורכבות בפועל עם הארכיטקט.'])
    add_phase_box_red(doc, '\u05e9\u05dc\u05d1 4 \u2014 \u05e1\u05e0\u05db\u05e8\u05d5\u05df \u05de\u05dc\u05d0 \u05d5-FHIR (6-12 \u05d7\u05d5\u05d3\u05e9\u05d9\u05dd)', [
        '\u05e1\u05e0\u05db\u05e8\u05d5\u05df \u05d3\u05d5-\u05db\u05d9\u05d5\u05d5\u05e0\u05d9 \u05de\u05dc\u05d0 \u05db\u05d5\u05dc\u05dc \u05ea\u05d5\u05e6\u05d0\u05d5\u05ea \u05de\u05e2\u05d1\u05d3\u05d4, \u05d2\u05d5\u05e8\u05de\u05d9 \u05e1\u05d9\u05db\u05d5\u05df, \u05d4\u05d9\u05e1\u05d8\u05d5\u05e8\u05d9\u05d9\u05ea \u05d1\u05d9\u05e7\u05d5\u05e8\u05d9\u05dd, \u05d5\u05de\u05e2\u05d1\u05e8 \u05dc\u05ea\u05e7\u05df FHIR \u05db\u05e9\u05d9\u05d4\u05d9\u05d4 \u05d6\u05de\u05d9\u05df (2027).',
    ])
    add_heading_rtl(doc, '12.1 \u05e4\u05e8\u05d9\u05d8\u05d9 \u05e4\u05e2\u05d5\u05dc\u05d4', level=2, color=BRAND_BLUE)
    add_bullet_list(doc, [
        '**1. \u05e4\u05d2\u05d9\u05e9\u05ea \u05d1\u05d9\u05e8\u05d5\u05e8 \u05de\u05d7\u05d9\u05e6\u05d4:** \u05dc\u05e7\u05d1\u05d5\u05e2 \u05e4\u05d2\u05d9\u05e9\u05d4 \u05e2\u05dd \u05e2\u05e0\u05ea \u05de\u05d9\u05e8\u05d5\u05df \u05d5\u05e6\u05d5\u05d5\u05ea \u05d4\u05d0\u05d3\u05e8\u05d9\u05db\u05dc\u05d5\u05ea \u05dc\u05d1\u05d9\u05e8\u05d5\u05e8 \u05db\u05dc\u05dc\u05d9 \u05d4\u05de\u05d7\u05d9\u05e6\u05d4 \u05d5\u05d4\u05e9\u05dc\u05db\u05d5\u05ea\u05d9\u05d4 \u05e2\u05dc \u05d4-API \u2014 \u05d0\u05d7\u05e8\u05d0\u05d9: \u05d3\u05e7\u05dc \u05ea\u05dc\u05d9\u05d0\u05d6, \u05e2\u05d3 \u05e9\u05d1\u05d5\u05e2\u05d9\u05d9\u05dd',
        '**2. \u05de\u05e1\u05de\u05da \u05d8\u05db\u05e0\u05d9 \u05dc\u05d0\u05d3\u05e8\u05d9\u05db\u05dc\u05d5\u05ea:** \u05dc\u05d4\u05e2\u05d1\u05d9\u05e8 \u05de\u05e1\u05de\u05da \u05d6\u05d4 (v2) \u05dc\u05e9\u05dc\u05d5\u05de\u05d9 \u05dc\u05d1\u05d7\u05d9\u05e0\u05ea \u05e6\u05d5\u05d5\u05ea \u05d4\u05d0\u05d3\u05e8\u05d9\u05db\u05dc\u05d5\u05ea \u2014 \u05d0\u05d7\u05e8\u05d0\u05d9: \u05d3\u05f4\u05e8 \u05e8\u05d0\u05de\u05d9 \u05d7\u05f3\u05d5\u05e8\u05d9, \u05de\u05d9\u05d9\u05d3\u05d9',
        '**3. \u05e1\u05d1\u05d9\u05d1\u05ea TEST:** \u05dc\u05e7\u05d1\u05dc \u05d2\u05d9\u05e9\u05d4 \u05dc\u05e1\u05d1\u05d9\u05d1\u05ea \u05d1\u05d3\u05d9\u05e7\u05d5\u05ea \u05e9\u05dc \u05e7\u05dc\u05d9\u05e7\u05e1 \u2014 \u05d0\u05d7\u05e8\u05d0\u05d9: \u05e9\u05dc\u05d5\u05de\u05d9 (\u05de\u05db\u05d1\u05d9), \u05e2\u05d3 3 \u05e9\u05d1\u05d5\u05e2\u05d5\u05ea',
        '**4. PoC \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5:** \u05dc\u05e4\u05ea\u05d7 Proof of Concept \u05dc\u05d4\u05e2\u05dc\u05d0\u05ea \u05ea\u05e9\u05d5\u05d1\u05ea \u05d9\u05d5\u05e2\u05e5 \u05d1-TEST \u2014 \u05d0\u05d7\u05e8\u05d0\u05d9: \u05e6\u05d5\u05d5\u05ea \u05ea\u05dc\u05d9\u05d0\u05d6, \u05d7\u05d5\u05d3\u05e9 \u05de\u05e7\u05d1\u05dc\u05ea \u05d2\u05d9\u05e9\u05d4',
        '**5. \u05d1\u05d7\u05d9\u05e0\u05ea \u05d1\u05ea \u05e7\u05d5\u05dc:** \u05dc\u05d1\u05d3\u05d5\u05e7 \u05d4\u05d0\u05dd \u05e0\u05d9\u05ea\u05df \u05dc\u05d4\u05e9\u05ea\u05de\u05e9 \u05d1\u05ea\u05e9\u05ea\u05d9\u05ea \u05d4-API \u05e9\u05dc \u05e4\u05e8\u05d5\u05d9\u05e7\u05d8 \u05d1\u05ea \u05e7\u05d5\u05dc \u2014 \u05d0\u05d7\u05e8\u05d0\u05d9: \u05e9\u05dc\u05d5\u05de\u05d9, \u05e9\u05d1\u05d5\u05e2\u05d9\u05d9\u05dd',
        '**6. \u05d9\u05e9\u05d9\u05d1\u05ea \u05de\u05e2\u05e7\u05d1:** \u05dc\u05ea\u05d0\u05dd \u05d9\u05e9\u05d9\u05d1\u05ea \u05de\u05e2\u05e7\u05d1 \u05dc\u05d0\u05d7\u05e8 \u05e7\u05d1\u05dc\u05ea \u05ea\u05e9\u05d5\u05d1\u05d4 \u05de\u05e6\u05d5\u05d5\u05ea \u05d4\u05d0\u05d3\u05e8\u05d9\u05db\u05dc\u05d5\u05ea \u2014 \u05d0\u05d7\u05e8\u05d0\u05d9: \u05d3\u05f4\u05e8 \u05e8\u05d0\u05de\u05d9 + \u05e9\u05dc\u05d5\u05de\u05d9, 3-4 \u05e9\u05d1\u05d5\u05e2\u05d5\u05ea',
    ])
    add_heading_rtl(doc, '12.2 עלויות והיבטים כלכליים', level=2, color=BRAND_BLUE)
    add_quote_box(doc, 'יהודה (מכבי) — מהישיבה:', '"מבחינה כלכלית — הרי זה פרויקט, זה צורך שלכם. תליאז, אנחנו לא היינו יוצאים לפיתוח לולא לא היה לכם את הקושי הזה עם הפסיכיאטרים, ולכן העלות צריכה להיות שלכם."')
    add_quote_box(doc, 'דקל תליאז (מנכ"ל תליאז) — מהישיבה:', '"אני לא חושב, אבל בוא לא נלאה את שלומי פה באירוע הזה. אני לא מסכים."')
    add_quote_box(doc, 'שלומי (מכבי) — מהישיבה:', '"אני מסכים, באמת תוציאו אותי מהתימונה כרגע. אני יכול לבוא עם הערכות ולהסביר מה המשמעות [...] שנשים ארכיטקט חינם, בסדר, כרגע לא עולה כסף."')
    add_decision_box(doc, 'סיכום כלכלי:', ['קיימת מחלוקת בין יהודה (מכבי) לדקל (תליאז) לגבי נשיאת עלויות הפיתוח. שלומי ביקש לדחות את הדיון הכלכלי ולהתחיל בשלב ההערכה הטכנית ללא עלות. ההחלטה הכלכלית תתקבל לאחר שהארכיטקט יציג את ההערכה ליהודה.'])

    # ==================== FOOTER ====================
    doc.add_page_break()
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_paragraph(p1)
    pPr = p1._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="6" w:space="4" w:color="888888"/></w:pBdr>')
    pPr.append(pBdr)
    add_hebrew_run(p1, '\u05de\u05e1\u05de\u05da \u05d6\u05d4 \u05d4\u05d5\u05e4\u05e7 \u05e2\u05dc \u05d1\u05e1\u05d9\u05e1:', bold=True, size=Pt(9), color=LIGHT_GRAY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_rtl_paragraph(p2)
    add_hebrew_run(p2, '(1) \u05d9\u05e9\u05d9\u05d1\u05ea \u05ea\u05d9\u05d0\u05d5\u05dd \u05ea\u05dc\u05d9\u05d0\u05d6-\u05de\u05db\u05d1\u05d9, 19 \u05d1\u05de\u05e8\u05e5 2026', size=Pt(9), color=LIGHT_GRAY)
    p2.add_run('\n')
    add_hebrew_run(p2, '(2) \u05e9\u05e8\u05e9\u05e8\u05ea \u05d0\u05d9\u05de\u05d9\u05d9\u05dc \u05d1\u05e0\u05d5\u05e9\u05d0 \u05d0\u05d9\u05e0\u05d8\u05d2\u05e8\u05e6\u05d9\u05d9\u05ea \u05e7\u05dc\u05d9\u05e7\u05e1, \u05de\u05e8\u05e5 2026', size=Pt(9), color=LIGHT_GRAY)
    p2.add_run('\n')
    add_hebrew_run(p2, '(3) \u05d4\u05e7\u05dc\u05d8\u05ea \u05de\u05e1\u05da \u05e9\u05dc \u05d0\u05dc\u05d9\u05f3\u05d0\u05e1 \u05d2\u05f3\u05d4\u05d2\u05f3\u05d0\u05d4, 22 \u05d1\u05e0\u05d5\u05d1\u05de\u05d1\u05e8 2024 \u2014 \u05d6\u05e8\u05d9\u05de\u05ea \u05e2\u05d1\u05d5\u05d3\u05d4 \u05de\u05dc\u05d0\u05d4 \u05d1\u05e7\u05dc\u05d9\u05e7\u05e1 EHR', size=Pt(9), color=LIGHT_GRAY)

    return doc

def main():
    print("Building Hebrew RTL Word document...")
    doc = build_document()

    print(f"Saving to: {OUTPUT_DOCX}")
    doc.save(str(OUTPUT_DOCX))

    size_mb = os.path.getsize(OUTPUT_DOCX) / (1024 * 1024)
    print(f"Word document saved: {OUTPUT_DOCX}")
    print(f"File size: {size_mb:.1f} MB")


if __name__ == '__main__':
    main()
